import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
from datetime import datetime, timedelta
import io
import base64
from scipy.stats import chi2_contingency
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

# =====================================================
# CONFIGURACIÓN DE LA PÁGINA STREAMLIT
# =====================================================
st.set_page_config(
    page_title="Análisis de Campaña Electoral",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =====================================================
# CONFIGURACIÓN VISUAL Y FORMATO APA
# =====================================================

# Configuración para cumplir con estándares APA
plt.style.use('default')  # Estilo base limpio
plt.rcParams.update({
    # FUENTES (APA recomienda Times New Roman 12pt, pero usaremos serif similar)
    'font.family': 'serif',
    'font.serif': ['Times New Roman', 'DejaVu Serif', 'serif'],
    'font.size': 12,  # Tamaño base APA
    'axes.titlesize': 14,  # Títulos de ejes
    'axes.labelsize': 12,  # Etiquetas de ejes
    'xtick.labelsize': 11,  # Etiquetas del eje X
    'ytick.labelsize': 11,  # Etiquetas del eje Y
    'legend.fontsize': 11,  # Leyenda
    'figure.titlesize': 16,  # Título principal
    
    # COLORES Y ESTILOS (APA prefiere escala de grises, pero permitimos color moderado)
    'axes.spines.top': False,     # Sin borde superior
    'axes.spines.right': False,   # Sin borde derecho
    'axes.spines.left': True,     # Borde izquierdo
    'axes.spines.bottom': True,   # Borde inferior
    'axes.linewidth': 1.0,        # Grosor de bordes
    'axes.edgecolor': 'black',    # Color de bordes
    'axes.facecolor': 'white',    # Fondo blanco
    'figure.facecolor': 'white',  # Fondo de figura
    
    # GRILLAS (opcional, ajustar según preferencia)
    'axes.grid': True,            # Activar grillas
    'grid.alpha': 0.3,            # Transparencia de grillas
    'grid.linewidth': 0.5,        # Grosor de grillas
    'grid.color': 'gray',         # Color de grillas
    
    # TAMAÑOS DE FIGURA (ajustar según necesidad)
    'figure.figsize': (10, 6),    # Tamaño por defecto
    'figure.dpi': 300,            # Resolución alta para publicación
    'savefig.dpi': 300,           # DPI al guardar
    'savefig.bbox': 'tight',      # Recorte ajustado
    'savefig.facecolor': 'white', # Fondo blanco al guardar
})

# Paleta de colores profesional
COLORES_PRINCIPALES = [
    '#2E86AB',  # Azul profesional
    '#A23B72',  # Magenta
    '#F18F01',  # Naranja
    '#C73E1D',  # Rojo
    '#8B5A3C',  # Marrón
    '#4A4A4A',  # Gris oscuro
    '#7B68EE',  # Azul lavanda
    '#32CD32',  # Verde lima
]

# Configurar seaborn para que coincida
sns.set_palette(COLORES_PRINCIPALES)

# =====================================================
# CONFIGURACIÓN DE LA PÁGINA  
# =====================================================

# La configuración se hace dentro de main()

# CSS personalizado se aplica dentro de main()

# =====================================================
# FUNCIONES AUXILIARES
# =====================================================

def obtener_variables_principales(dummy_cols):
    """Obtiene la lista de variables principales (sin las subcategorías)"""
    variables = set()
    for col in dummy_cols:
        if '__' in col:
            variable_principal = col.split('__')[0]
            variables.add(variable_principal)
    return sorted(list(variables))

def obtener_categorias_de_variable(dummy_cols, variable_principal):
    """Obtiene las categorías de una variable específica"""
    categorias = []
    for col in dummy_cols:
        if '__' in col and col.split('__')[0] == variable_principal:
            categoria = col.split('__')[1]
            categorias.append((col, categoria))
    return categorias

def filtrar_datos_por_seleccion(df, dummy_cols, variable_seleccionada=None, categorias_seleccionadas=None):
    """Filtra el DataFrame según la selección de variable/categorías (múltiples)"""
    if variable_seleccionada == "Todas las variables" or variable_seleccionada is None:
        return df, dummy_cols
    
    # Filtrar columnas dummy de la variable seleccionada
    if not categorias_seleccionadas or "Todas las categorías" in categorias_seleccionadas:
        cols_filtradas = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
        df_filtrado = df.copy()
    else:
        # Selecciones múltiples específicas
        cols_filtradas = []
        condiciones = []
        
        for categoria in categorias_seleccionadas:
            col_especifica = f"{variable_seleccionada}__{categoria}"
            if col_especifica in df.columns:
                cols_filtradas.append(col_especifica)
                condiciones.append(df[col_especifica] == 1)
        
        # Combinar condiciones con OR (al menos una categoría debe ser verdadera)
        if condiciones:
            condicion_final = condiciones[0]
            for condicion in condiciones[1:]:
                condicion_final = condicion_final | condicion
            df_filtrado = df[condicion_final].copy()
        else:
            df_filtrado = df.copy()
    
    return df_filtrado, cols_filtradas

def aplicar_formato_apa_dataframe(df, titulo="Tabla", incluir_nota=True):
    """Aplica formato APA a un DataFrame para visualización en Streamlit"""
    df_formateado = df.copy()
    
    # Redondear números decimales a 2 lugares
    for col in df_formateado.select_dtypes(include=[np.number]).columns:
        df_formateado[col] = df_formateado[col].round(2)
    
    return df_formateado

def mostrar_tabla_con_formato(df, titulo, formato_apa=False):
    """Muestra una tabla con el formato seleccionado"""
    if formato_apa:
        df_mostrar = aplicar_formato_apa_dataframe(df, titulo)
        st.markdown(f"**{titulo}**")
        st.dataframe(df_mostrar, use_container_width=True)
        if len(df) > 0:
            st.caption("*Nota*. Los valores están redondeados a dos decimales según estándares APA.")
    else:
        st.subheader(titulo)
        st.dataframe(df, use_container_width=True)

@st.cache_data
def cargar_datos():
    """Carga y procesa los datos del archivo Excel"""
    try:
        df = pd.read_excel("recodificado.xlsx")
        
        # Procesar fechas
        if 'Fecha' in df.columns:
            meses = {
                'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4, 'mayo': 5, 'junio': 6,
                'julio': 7, 'agosto': 8, 'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12
            }
            
            def convertir_fecha(fecha_str):
                try:
                    if pd.isna(fecha_str):
                        return None
                    partes = str(fecha_str).lower().strip().split(' de ')
                    if len(partes) == 2:
                        dia = int(partes[0])
                        mes = meses.get(partes[1], 1)
                        return datetime(2025, mes, dia)
                except:
                    pass
                return None
            
            df['Fecha_convertida'] = df['Fecha'].apply(convertir_fecha)
        
        # Identificar columnas dummy
        dummy_cols = [col for col in df.columns if '__' in col]
        
        return df, dummy_cols
    except Exception as e:
        st.error(f"Error al cargar datos: {e}")
        return None, []

def crear_ranking_por_variable(df, dummy_cols, variable_seleccionada=None, n_top=10, formato_apa=False):
    """Crea ranking de categorías dentro de una variable específica o de todas"""
    try:
        if variable_seleccionada and variable_seleccionada != "Todas las variables":
            # Filtrar solo las columnas de la variable seleccionada
            cols_variable = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
        else:
            cols_variable = dummy_cols
        
        # Calcular sumas
        sumas = {}
        for col in cols_variable:
            if col in df.columns:
                suma = df[col].sum()
                if suma > 0:  # Solo incluir categorías con al menos 1 uso
                    if '__' in col:
                        variable = col.split('__')[0].replace('_', ' ').title()
                        categoria = col.split('__')[1].replace('_', ' ').title()
                        nombre_completo = f"{variable} - {categoria}"
                    else:
                        nombre_completo = col
                    sumas[nombre_completo] = suma
        
        # Crear DataFrame y ordenar
        if sumas:
            df_ranking = pd.DataFrame(list(sumas.items()), columns=['Categoría', 'Frecuencia'])
            df_ranking = df_ranking.sort_values('Frecuencia', ascending=False).head(n_top)
            df_ranking = df_ranking.reset_index(drop=True)
            df_ranking.index += 1  # Comenzar numeración en 1
            
            # Agregar porcentaje
            total = df_ranking['Frecuencia'].sum()
            df_ranking['Porcentaje'] = (df_ranking['Frecuencia'] / len(df) * 100).round(2)
            
            return df_ranking
        else:
            return pd.DataFrame(columns=['Categoría', 'Frecuencia', 'Porcentaje'])
            
    except Exception as e:
        st.error(f"Error al crear ranking: {e}")
        return pd.DataFrame(columns=['Categoría', 'Frecuencia', 'Porcentaje'])

def crear_ranking_por_candidato_y_total(df, dummy_cols, variable_seleccionada=None, n_top=10, formato_apa=False):
    """Crea ranking de categorías por candidato y total general"""
    try:
        if variable_seleccionada and variable_seleccionada != "Todas las variables":
            # Filtrar solo las columnas de la variable seleccionada
            cols_variable = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
        else:
            cols_variable = dummy_cols
        
        # Obtener candidatos únicos
        candidatos = df['Candidato'].unique() if 'Candidato' in df.columns else ['Sin candidato']
        
        # Crear diccionario para almacenar resultados
        resultados = {}
        
        # Calcular para cada candidato
        for candidato in candidatos:
            df_candidato = df[df['Candidato'] == candidato] if 'Candidato' in df.columns else df
            sumas_candidato = {}
            
            for col in cols_variable:
                if col in df_candidato.columns:
                    suma = df_candidato[col].sum()
                    if suma > 0:
                        if '__' in col:
                            variable = col.split('__')[0].replace('_', ' ').title()
                            categoria = col.split('__')[1].replace('_', ' ').title()
                            nombre_completo = f"{variable} - {categoria}"
                        else:
                            nombre_completo = col
                        sumas_candidato[nombre_completo] = suma
            
            # Crear DataFrame para este candidato
            if sumas_candidato:
                df_cand = pd.DataFrame(list(sumas_candidato.items()), columns=['Categoría', f'{candidato}'])
                df_cand = df_cand.sort_values(f'{candidato}', ascending=False)
                resultados[candidato] = df_cand
        
        # Calcular totales generales
        sumas_total = {}
        for col in cols_variable:
            if col in df.columns:
                suma = df[col].sum()
                if suma > 0:
                    if '__' in col:
                        variable = col.split('__')[0].replace('_', ' ').title()
                        categoria = col.split('__')[1].replace('_', ' ').title()
                        nombre_completo = f"{variable} - {categoria}"
                    else:
                        nombre_completo = col
                    sumas_total[nombre_completo] = suma
        
        # Crear DataFrame consolidado
        if sumas_total:
            # Comenzar con totales
            df_consolidado = pd.DataFrame(list(sumas_total.items()), columns=['Categoría', 'Total'])
            df_consolidado = df_consolidado.sort_values('Total', ascending=False).head(n_top)
            
            # Agregar columnas por candidato
            for candidato in candidatos:
                df_consolidado[candidato] = 0
                if candidato in resultados:
                    df_cand = resultados[candidato]
                    for idx, row in df_consolidado.iterrows():
                        categoria = row['Categoría']
                        match = df_cand[df_cand['Categoría'] == categoria]
                        if not match.empty:
                            df_consolidado.at[idx, candidato] = match.iloc[0][candidato]
            
            # Calcular porcentajes
            total_general = df_consolidado['Total'].sum()
            df_consolidado['Porcentaje'] = (df_consolidado['Total'] / len(df) * 100).round(2)
            
            # Resetear índice y empezar en 1
            df_consolidado = df_consolidado.reset_index(drop=True)
            df_consolidado.index += 1
            
            return df_consolidado
        else:
            columns = ['Categoría', 'Total'] + list(candidatos) + ['Porcentaje']
            return pd.DataFrame(columns=columns)
            
    except Exception as e:
        st.error(f"Error al crear ranking por candidato: {e}")
        return pd.DataFrame()

def crear_tabla_cruzada(df, var1, var2, formato_apa=False):
    """Crea una tabla cruzada entre dos variables"""
    try:
        # Crear tabla de contingencia
        tabla = pd.crosstab(df[var1], df[var2], margins=True)
        
        # Crear tabla de porcentajes
        tabla_pct = pd.crosstab(df[var1], df[var2], normalize='all') * 100
        tabla_pct = tabla_pct.round(2)
        
        # Estadística Chi-cuadrado
        chi2_stat, p_value = None, None
        try:
            if tabla.shape[0] > 1 and tabla.shape[1] > 1:
                # Excluir márgenes para el test
                tabla_test = tabla.iloc[:-1, :-1]
                if tabla_test.sum().sum() > 0:
                    chi2_stat, p_value, _, _ = chi2_contingency(tabla_test)
        except:
            pass
        
        return tabla, tabla_pct, chi2_stat, p_value
    except Exception as e:
        st.error(f"Error al crear tabla cruzada: {e}")
        return None, None, None, None

def exportar_a_excel(dataframes_dict, nombre_archivo):
    """Exporta múltiples DataFrames a un archivo Excel"""
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for sheet_name, df in dataframes_dict.items():
            df.to_excel(writer, sheet_name=sheet_name[:31], index=False)  # Límite de 31 caracteres para nombres de hojas
    
    output.seek(0)
    return output

def analisis_evolucion_temporal(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de evolución temporal de estrategias"""
    if 'Fecha_convertida' not in df.columns:
        return None, []
    
    df_temporal = df.dropna(subset=['Fecha_convertida'])
    
    if variable_seleccionada and variable_seleccionada != "Todas las variables":
        # Filtrar columnas de la variable seleccionada
        estrategias_clave = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada][:6]
    else:
        # Estrategias clave para seguimiento temporal (todas las variables)
        estrategias_clave = [
            col for col in dummy_cols 
            if any(palabra in col.lower() for palabra in ['meme', 'logotipo', 'testimonio', 'plain_folks', 'orquestacion'])
        ][:6]
    
    if not estrategias_clave:
        return None, []
    
    # Agrupar por fecha
    df_agrupado = df_temporal.groupby('Fecha_convertida')[estrategias_clave].sum().reset_index()
    
    return df_agrupado, estrategias_clave

def analisis_propaganda_candidatos(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de técnicas de propaganda por candidato"""
    if variable_seleccionada and variable_seleccionada != "Todas las variables":
        propaganda_cols = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
    else:
        propaganda_cols = [col for col in dummy_cols if 'institute' in col.lower() or 'propaganda' in col.lower()]
    
    if not propaganda_cols or 'Candidato' not in df.columns:
        return None
    
    candidatos = df['Candidato'].unique()
    datos_propaganda = []
    
    for candidato in candidatos:
        df_cand = df[df['Candidato'] == candidato]
        total_posts = len(df_cand)
        
        for col in propaganda_cols:
            if col in df.columns:
                usos = df_cand[col].sum()
                porcentaje = (usos / total_posts) * 100 if total_posts > 0 else 0
                
                if '__' in col:
                    variable = col.split('__')[0].replace('_', ' ').title()
                    tecnica = col.split('__')[1].replace('_', ' ').title()
                    nombre_completo = f"{variable} - {tecnica}"
                else:
                    nombre_completo = col
                
                datos_propaganda.append({
                    'Candidato': candidato,
                    'Técnica': nombre_completo,
                    'Usos': usos,
                    'Porcentaje': porcentaje
                })
    
    return pd.DataFrame(datos_propaganda)

def analisis_plain_folks(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis detallado de la estrategia Plain-folks"""
    if variable_seleccionada and variable_seleccionada != "Todas las variables":
        plain_folks_cols = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada and 'plain' in col.lower()]
    else:
        plain_folks_cols = [col for col in dummy_cols if 'plain' in col.lower() or 'pueblo' in col.lower()]
    
    if not plain_folks_cols:
        return None
    
    # Filtrar posts que usan Plain-folks
    df_plain = df[df[plain_folks_cols].sum(axis=1) > 0].copy()
    
    resultados = {}
    
    # Por candidato
    if 'Candidato' in df.columns and len(df_plain) > 0:
        candidatos_stats = []
        for candidato in df['Candidato'].unique():
            df_cand = df[df['Candidato'] == candidato]
            df_plain_cand = df_plain[df_plain['Candidato'] == candidato]
            
            total_posts = len(df_cand)
            plain_posts = len(df_plain_cand)
            porcentaje = (plain_posts / total_posts) * 100 if total_posts > 0 else 0
            
            candidatos_stats.append({
                'Candidato': candidato,
                'Total_Posts': total_posts,
                'Plain_Folks_Posts': plain_posts,
                'Porcentaje': porcentaje
            })
        
        resultados['candidatos'] = pd.DataFrame(candidatos_stats)
    
    return resultados

# =====================================================
# NUEVAS FUNCIONES DE ANÁLISIS AVANZADO
# =====================================================

def analisis_plain_folks_por_contexto(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis detallado de Plain-folks según contexto y campaña"""
    resultados = {}
    
    # Buscar columnas de Plain-folks
    plain_folks_cols = [col for col in dummy_cols if 'plain_folks' in col.lower() or 'plain_folk' in col.lower()]
    
    if not plain_folks_cols:
        return None
    
    # Análisis por contexto y candidato
    if 'Contexto_de_la_imagen' in df.columns and 'Candidato' in df.columns:
        contextos = df['Contexto_de_la_imagen'].dropna().unique()
        candidatos = df['Candidato'].dropna().unique()
        
        # Crear tabla de Plain-folks por contexto y candidato
        tabla_contexto = []
        
        for contexto in contextos:
            df_contexto = df[df['Contexto_de_la_imagen'] == contexto]
            
            # Calcular totales por candidato
            for candidato in candidatos:
                df_cand_contexto = df_contexto[df_contexto['Candidato'] == candidato]
                
                if len(df_cand_contexto) > 0:
                    # Contar publicaciones con Plain-folks
                    plain_folks_count = 0
                    for col in plain_folks_cols:
                        if col in df_cand_contexto.columns:
                            plain_folks_count += df_cand_contexto[col].sum()
                    
                    total_publicaciones = len(df_cand_contexto)
                    porcentaje = (plain_folks_count / total_publicaciones * 100) if total_publicaciones > 0 else 0
                    
                    tabla_contexto.append({
                        'Contexto': contexto,
                        'Campaña': candidato,
                        f'{candidato}: publicaciones con Plain-folks (%)': f"{plain_folks_count} ({porcentaje:.1f} %)",
                        'Total_publicaciones': total_publicaciones,
                        'Plain_folks_count': plain_folks_count
                    })
        
        if tabla_contexto:
            df_contexto_resultado = pd.DataFrame(tabla_contexto)
            
            # Reorganizar para mostrar como en la imagen
            tabla_pivot = []
            for contexto in contextos:
                fila = {'Contexto': contexto, 'Campaña': 'Votantes' if 'Vía pública' in contexto else 'Familiares'}
                
                for candidato in candidatos:
                    datos_contexto = df_contexto_resultado[
                        (df_contexto_resultado['Contexto'] == contexto)
                    ]
                    
                    for _, row in datos_contexto.iterrows():
                        if candidato in row[f'{candidato}: publicaciones con Plain-folks (%)']:
                            fila[f'{candidato}: publicaciones con Plain-folks (%)'] = row[f'{candidato}: publicaciones con Plain-folks (%)']
                
                tabla_pivot.append(fila)
            
            resultados['plain_folks_contexto'] = pd.DataFrame(tabla_pivot)
    
    return resultados

def analisis_distribucion_propaganda_ipa(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de distribución general de recursos de propaganda según IPA"""
    resultados = {}
    
    # Recursos de propaganda IPA
    recursos_ipa = [
        'name_calling', 'glittering_generalities', 'transfer', 
        'testimonial', 'plain_folks', 'card_stacking', 'bandwagon'
    ]
    
    if 'Candidato' in df.columns:
        candidatos = df['Candidato'].dropna().unique()
        
        tabla_recursos = []
        
        for i, recurso in enumerate(recursos_ipa, 1):
            fila = {'Recurso de propaganda (IPA)': f"{i}. {recurso.replace('_', '-').title()}"}
            
            for candidato in candidatos:
                df_candidato = df[df['Candidato'] == candidato]
                
                # Buscar columnas que contengan este recurso
                recurso_cols = [col for col in dummy_cols if recurso.lower() in col.lower()]
                
                total_recurso = 0
                for col in recurso_cols:
                    if col in df_candidato.columns:
                        total_recurso += df_candidato[col].sum()
                
                total_publicaciones = len(df_candidato)
                porcentaje = (total_recurso / total_publicaciones * 100) if total_publicaciones > 0 else 0
                
                fila[f'{candidato}: Nº de publicaciones'] = total_recurso
                fila[f'{candidato}: %'] = f"{porcentaje:.1f}"
            
            tabla_recursos.append(fila)
        
        resultados['distribucion_ipa'] = pd.DataFrame(tabla_recursos)
    
    return resultados

def analisis_cruce_reglas_contexto(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de cruce entre reglas de propaganda y contexto de imagen"""
    resultados = {}
    
    # Buscar variables de reglas dominantes y contexto
    if 'Contexto_de_la_imagen' in df.columns:
        contextos = df['Contexto_de_la_imagen'].dropna().unique()
        
        # Buscar columnas de reglas de propaganda
        reglas_cols = [col for col in dummy_cols if any(regla in col.lower() for regla in 
                      ['name_calling', 'glittering', 'transfer', 'testimonial', 'plain_folks', 'card_stacking', 'bandwagon'])]
        
        tabla_cruce = []
        
        for contexto in contextos:
            df_contexto = df[df['Contexto_de_la_imagen'] == contexto]
            
            # Encontrar la regla dominante en este contexto
            regla_dominante = None
            max_count = 0
            
            for col in reglas_cols:
                if col in df_contexto.columns:
                    count = df_contexto[col].sum()
                    if count > max_count:
                        max_count = count
                        regla_dominante = col.split('__')[-1] if '__' in col else col
            
            # Determinar campaña dominante
            campana_dominante = "Votantes"  # Por defecto
            candidato_dominante = "Ambos"
            
            if 'Candidato' in df_contexto.columns:
                candidatos_count = df_contexto['Candidato'].value_counts()
                if len(candidatos_count) > 0:
                    candidato_dominante = candidatos_count.index[0]
            
            # Mapear contexto a campaña
            if 'personal' in contexto.lower():
                campana_dominante = "Familiares"
            elif 'profesional' in contexto.lower():
                campana_dominante = "Ninguno"
            elif 'pública' in contexto.lower() or 'publica' in contexto.lower():
                campana_dominante = "Votantes"
            
            tabla_cruce.append({
                'Contexto de la imagen': contexto,
                'Regla dominante (Domenach)': regla_dominante.replace('_', ' ').title() if regla_dominante else "N/A",
                'Nº publicaciones': len(df_contexto),
                'Candidato dominante': candidato_dominante
            })
        
        resultados['cruce_reglas_contexto'] = pd.DataFrame(tabla_cruce)
    
    return resultados

def analisis_aparicion_lider(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de aparición del líder según contexto y campaña"""
    resultados = {}
    
    # Buscar columnas relacionadas con aparición del líder
    aparicion_cols = [col for col in dummy_cols if 'aparicion' in col.lower() or 'lider' in col.lower() or 'acompañado' in col.lower()]
    
    if 'Contexto_de_la_imagen' in df.columns and aparicion_cols:
        contextos = df['Contexto_de_la_imagen'].dropna().unique()
        
        tabla_aparicion = []
        
        for contexto in contextos:
            df_contexto = df[df['Contexto_de_la_imagen'] == contexto]
            
            # Determinar si va acompañado
            acompanado = "Ninguno"
            if 'personal' in contexto.lower():
                acompanado = "Familiares"
            elif 'pública' in contexto.lower() or 'publica' in contexto.lower():
                acompanado = "Votantes"
            
            # Calcular presencia del líder
            presencia_lider = 0
            total_publicaciones = len(df_contexto)
            
            for col in aparicion_cols:
                if col in df_contexto.columns:
                    presencia_lider += df_contexto[col].sum()
            
            porcentaje_presencia = (presencia_lider / total_publicaciones * 100) if total_publicaciones > 0 else 0
            
            tabla_aparicion.append({
                'Contexto': contexto,
                '¿Va acompañado?': acompanado,
                'Presencia del líder (%)': f"{porcentaje_presencia:.0f}%",
                'Total': total_publicaciones
            })
        
        resultados['aparicion_lider'] = pd.DataFrame(tabla_aparicion)
    
    return resultados

def generar_tabla_contingencia_avanzada(df, var1, var2, incluir_porcentajes=True):
    """Genera tabla de contingencia avanzada con múltiples estadísticos"""
    try:
        # Crear tabla de contingencia básica
        tabla = pd.crosstab(df[var1], df[var2], margins=True)
        
        resultados = {'tabla_frecuencias': tabla}
        
        if incluir_porcentajes:
            # Tabla de porcentajes por filas
            tabla_pct_filas = pd.crosstab(df[var1], df[var2], normalize='index') * 100
            tabla_pct_filas = tabla_pct_filas.round(1)
            
            # Tabla de porcentajes por columnas
            tabla_pct_cols = pd.crosstab(df[var1], df[var2], normalize='columns') * 100
            tabla_pct_cols = tabla_pct_cols.round(1)
            
            # Tabla de porcentajes totales
            tabla_pct_total = pd.crosstab(df[var1], df[var2], normalize='all') * 100
            tabla_pct_total = tabla_pct_total.round(1)
            
            resultados.update({
                'tabla_porcentajes_filas': tabla_pct_filas,
                'tabla_porcentajes_columnas': tabla_pct_cols,
                'tabla_porcentajes_total': tabla_pct_total
            })
        
        # Prueba de chi-cuadrado si es posible
        try:
            from scipy.stats import chi2_contingency
            tabla_sin_margenes = tabla.iloc[:-1, :-1]  # Quitar totales
            chi2, p_valor, gl, esperados = chi2_contingency(tabla_sin_margenes)
            
            resultados['estadisticos'] = {
                'chi2': chi2,
                'p_valor': p_valor,
                'grados_libertad': gl,
                'significativo': p_valor < 0.05
            }
        except:
            pass
        
        return resultados
    
    except Exception as e:
        st.error(f"Error al generar tabla de contingencia: {str(e)}")
        return None

def crear_visualizacion_avanzada(df, tipo_analisis, **kwargs):
    """Crea visualizaciones avanzadas según el tipo de análisis"""
    figs = {}
    
    if tipo_analisis == "distribucion_ipa":
        # Gráfico de barras apiladas para recursos IPA
        if 'datos' in kwargs and not kwargs['datos'].empty:
            datos = kwargs['datos']
            
            # Preparar datos para gráfico
            candidatos = [col for col in datos.columns if 'Nº de publicaciones' in col]
            
            fig = go.Figure()
            
            for candidato_col in candidatos:
                candidato = candidato_col.split(':')[0]
                valores = datos[candidato_col].values
                recursos = datos['Recurso de propaganda (IPA)'].values
                
                fig.add_trace(go.Bar(
                    name=candidato,
                    x=recursos,
                    y=valores,
                    text=valores,
                    textposition='auto'
                ))
            
            fig.update_layout(
                title="Distribución de Recursos de Propaganda (IPA) por Candidato",
                xaxis_title="Recursos de Propaganda",
                yaxis_title="Número de Publicaciones",
                barmode='group',
                height=600
            )
            
            figs['distribucion_ipa'] = fig
    
    elif tipo_analisis == "plain_folks_contexto":
        # Gráfico de barras para Plain-folks por contexto
        if 'datos' in kwargs and not kwargs['datos'].empty:
            datos = kwargs['datos']
            
            fig = go.Figure()
            
            # Extraer datos para visualización
            contextos = datos['Contexto'].unique()
            candidatos_cols = [col for col in datos.columns if 'publicaciones con Plain-folks' in col]
            
            for col in candidatos_cols:
                candidato = col.split(':')[0]
                # Extraer números de los porcentajes
                valores = []
                for _, row in datos.iterrows():
                    texto = str(row[col])
                    if '(' in texto and ')' in texto:
                        porcentaje = texto.split('(')[1].split('%')[0]
                        try:
                            valores.append(float(porcentaje))
                        except:
                            valores.append(0)
                    else:
                        valores.append(0)
                
                fig.add_trace(go.Bar(
                    name=candidato,
                    x=datos['Contexto'],
                    y=valores,
                    text=[f"{v:.1f}%" for v in valores],
                    textposition='auto'
                ))
            
            fig.update_layout(
                title="Uso del recurso Plain-folks según contexto y campaña",
                xaxis_title="Contexto",
                yaxis_title="Porcentaje de publicaciones (%)",
                barmode='group',
                height=500
            )
            figs['plain_folks_contexto'] = fig
    
    return figs

# =====================================================
# NUEVAS FUNCIONES DE ANÁLISIS AVANZADO (CONTINUACIÓN) 
# =====================================================
# FUNCIONES AUXILIARES
# =====================================================

def obtener_variables_principales(dummy_cols):
    """Obtiene la lista de variables principales (sin las subcategorías)"""
    variables = set()
    for col in dummy_cols:
        if '__' in col:
            variable_principal = col.split('__')[0]
            variables.add(variable_principal)
    return sorted(list(variables))

def obtener_categorias_de_variable(dummy_cols, variable_principal):
    """Obtiene las categorías de una variable específica"""
    categorias = []
    for col in dummy_cols:
        if '__' in col and col.split('__')[0] == variable_principal:
            categoria = col.split('__')[1]
            categorias.append((col, categoria))
    return categorias

def filtrar_datos_por_seleccion(df, dummy_cols, variable_seleccionada=None, categorias_seleccionadas=None):
    """Filtra el DataFrame según la selección de variable/categorías (múltiples)"""
    if variable_seleccionada == "Todas las variables" or variable_seleccionada is None:
        return df, dummy_cols
    
    # Filtrar columnas dummy de la variable seleccionada
    if not categorias_seleccionadas or "Todas las categorías" in categorias_seleccionadas:
        cols_filtradas = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
        df_filtrado = df.copy()
    else:
        # Selecciones múltiples específicas
        cols_filtradas = []
        condiciones = []
        
        for categoria in categorias_seleccionadas:
            col_especifica = f"{variable_seleccionada}__{categoria}"
            if col_especifica in df.columns:
                cols_filtradas.append(col_especifica)
                condiciones.append(df[col_especifica] == 1)
        
        # Combinar condiciones con OR (al menos una categoría debe ser verdadera)
        if condiciones:
            condicion_final = condiciones[0]
            for condicion in condiciones[1:]:
                condicion_final = condicion_final | condicion
            df_filtrado = df[condicion_final].copy()
        else:
            df_filtrado = df.copy()
    
    return df_filtrado, cols_filtradas

def aplicar_formato_apa_dataframe(df, titulo="Tabla", incluir_nota=True):
    """Aplica formato APA a un DataFrame para visualización en Streamlit"""
    df_formateado = df.copy()
    
    # Redondear números decimales a 2 lugares
    for col in df_formateado.select_dtypes(include=[np.number]).columns:
        df_formateado[col] = df_formateado[col].round(2)
    
    return df_formateado

def mostrar_tabla_con_formato(df, titulo, formato_apa=False):
    """Muestra una tabla con el formato seleccionado"""
    if formato_apa:
        df_mostrar = aplicar_formato_apa_dataframe(df, titulo)
        st.markdown(f"**{titulo}**")
        st.dataframe(df_mostrar, use_container_width=True)
        if len(df) > 0:
            st.caption("*Nota*. Los valores están redondeados a dos decimales según estándares APA.")
    else:
        st.subheader(titulo)
        st.dataframe(df, use_container_width=True)

@st.cache_data
def cargar_datos():
    """Carga y procesa los datos del archivo Excel"""
    try:
        df = pd.read_excel("recodificado.xlsx")
        
        # Procesar fechas
        if 'Fecha' in df.columns:
            meses = {
                'enero': 1, 'febrero': 2, 'marzo': 3, 'abril': 4, 'mayo': 5, 'junio': 6,
                'julio': 7, 'agosto': 8, 'septiembre': 9, 'octubre': 10, 'noviembre': 11, 'diciembre': 12
            }
            
            def convertir_fecha(fecha_str):
                try:
                    if pd.isna(fecha_str):
                        return None
                    partes = str(fecha_str).lower().strip().split(' de ')
                    if len(partes) == 2:
                        dia = int(partes[0])
                        mes = meses.get(partes[1], 1)
                        return datetime(2025, mes, dia)
                except:
                    pass
                return None
            
            df['Fecha_convertida'] = df['Fecha'].apply(convertir_fecha)
        
        # Identificar columnas dummy
        dummy_cols = [col for col in df.columns if '__' in col]
        
        return df, dummy_cols
    except Exception as e:
        st.error(f"Error al cargar datos: {e}")
        return None, []

def crear_ranking_por_variable(df, dummy_cols, variable_seleccionada=None, n_top=10, formato_apa=False):
    """Crea ranking de categorías dentro de una variable específica o de todas"""
    try:
        if variable_seleccionada and variable_seleccionada != "Todas las variables":
            # Filtrar solo las columnas de la variable seleccionada
            cols_variable = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
        else:
            cols_variable = dummy_cols
        
        # Calcular sumas
        sumas = {}
        for col in cols_variable:
            if col in df.columns:
                suma = df[col].sum()
                if suma > 0:  # Solo incluir categorías con al menos 1 uso
                    if '__' in col:
                        variable = col.split('__')[0].replace('_', ' ').title()
                        categoria = col.split('__')[1].replace('_', ' ').title()
                        nombre_completo = f"{variable} - {categoria}"
                    else:
                        nombre_completo = col
                    sumas[nombre_completo] = suma
        
        # Crear DataFrame y ordenar
        if sumas:
            df_ranking = pd.DataFrame(list(sumas.items()), columns=['Categoría', 'Frecuencia'])
            df_ranking = df_ranking.sort_values('Frecuencia', ascending=False).head(n_top)
            df_ranking = df_ranking.reset_index(drop=True)
            df_ranking.index += 1  # Comenzar numeración en 1
            
            # Agregar porcentaje
            total = df_ranking['Frecuencia'].sum()
            df_ranking['Porcentaje'] = (df_ranking['Frecuencia'] / len(df) * 100).round(2)
            
            return df_ranking
        else:
            return pd.DataFrame(columns=['Categoría', 'Frecuencia', 'Porcentaje'])
            
    except Exception as e:
        st.error(f"Error al crear ranking: {e}")
        return pd.DataFrame(columns=['Categoría', 'Frecuencia', 'Porcentaje'])

def crear_ranking_por_candidato_y_total(df, dummy_cols, variable_seleccionada=None, n_top=10, formato_apa=False):
    """Crea ranking de categorías por candidato y total general"""
    try:
        if variable_seleccionada and variable_seleccionada != "Todas las variables":
            # Filtrar solo las columnas de la variable seleccionada
            cols_variable = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
        else:
            cols_variable = dummy_cols
        
        # Obtener candidatos únicos
        candidatos = df['Candidato'].unique() if 'Candidato' in df.columns else ['Sin candidato']
        
        # Crear diccionario para almacenar resultados
        resultados = {}
        
        # Calcular para cada candidato
        for candidato in candidatos:
            df_candidato = df[df['Candidato'] == candidato] if 'Candidato' in df.columns else df
            sumas_candidato = {}
            
            for col in cols_variable:
                if col in df_candidato.columns:
                    suma = df_candidato[col].sum()
                    if suma > 0:
                        if '__' in col:
                            variable = col.split('__')[0].replace('_', ' ').title()
                            categoria = col.split('__')[1].replace('_', ' ').title()
                            nombre_completo = f"{variable} - {categoria}"
                        else:
                            nombre_completo = col
                        sumas_candidato[nombre_completo] = suma
            
            # Crear DataFrame para este candidato
            if sumas_candidato:
                df_cand = pd.DataFrame(list(sumas_candidato.items()), columns=['Categoría', f'{candidato}'])
                df_cand = df_cand.sort_values(f'{candidato}', ascending=False)
                resultados[candidato] = df_cand
        
        # Calcular totales generales
        sumas_total = {}
        for col in cols_variable:
            if col in df.columns:
                suma = df[col].sum()
                if suma > 0:
                    if '__' in col:
                        variable = col.split('__')[0].replace('_', ' ').title()
                        categoria = col.split('__')[1].replace('_', ' ').title()
                        nombre_completo = f"{variable} - {categoria}"
                    else:
                        nombre_completo = col
                    sumas_total[nombre_completo] = suma
        
        # Crear DataFrame consolidado
        if sumas_total:
            # Comenzar con totales
            df_consolidado = pd.DataFrame(list(sumas_total.items()), columns=['Categoría', 'Total'])
            df_consolidado = df_consolidado.sort_values('Total', ascending=False).head(n_top)
            
            # Agregar columnas por candidato
            for candidato in candidatos:
                df_consolidado[candidato] = 0
                if candidato in resultados:
                    df_cand = resultados[candidato]
                    for idx, row in df_consolidado.iterrows():
                        categoria = row['Categoría']
                        match = df_cand[df_cand['Categoría'] == categoria]
                        if not match.empty:
                            df_consolidado.at[idx, candidato] = match.iloc[0][candidato]
            
            # Calcular porcentajes
            total_general = df_consolidado['Total'].sum()
            df_consolidado['Porcentaje'] = (df_consolidado['Total'] / len(df) * 100).round(2)
            
            # Resetear índice y empezar en 1
            df_consolidado = df_consolidado.reset_index(drop=True)
            df_consolidado.index += 1
            
            return df_consolidado
        else:
            columns = ['Categoría', 'Total'] + list(candidatos) + ['Porcentaje']
            return pd.DataFrame(columns=columns)
            
    except Exception as e:
        st.error(f"Error al crear ranking por candidato: {e}")
        return pd.DataFrame()

def crear_tabla_cruzada(df, var1, var2, formato_apa=False):
    """Crea una tabla cruzada entre dos variables"""
    try:
        # Crear tabla de contingencia
        tabla = pd.crosstab(df[var1], df[var2], margins=True)
        
        # Crear tabla de porcentajes
        tabla_pct = pd.crosstab(df[var1], df[var2], normalize='all') * 100
        tabla_pct = tabla_pct.round(2)
        
        # Estadística Chi-cuadrado
        chi2_stat, p_value = None, None
        try:
            if tabla.shape[0] > 1 and tabla.shape[1] > 1:
                # Excluir márgenes para el test
                tabla_test = tabla.iloc[:-1, :-1]
                if tabla_test.sum().sum() > 0:
                    chi2_stat, p_value, _, _ = chi2_contingency(tabla_test)
        except:
            pass
        
        return tabla, tabla_pct, chi2_stat, p_value
    except Exception as e:
        st.error(f"Error al crear tabla cruzada: {e}")
        return None, None, None, None

def exportar_a_excel(dataframes_dict, nombre_archivo):
    """Exporta múltiples DataFrames a un archivo Excel"""
    output = io.BytesIO()
    
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        for sheet_name, df in dataframes_dict.items():
            df.to_excel(writer, sheet_name=sheet_name[:31], index=False)  # Límite de 31 caracteres para nombres de hojas
    
    output.seek(0)
    return output

def analisis_evolucion_temporal(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de evolución temporal de estrategias"""
    if 'Fecha_convertida' not in df.columns:
        return None, []
    
    df_temporal = df.dropna(subset=['Fecha_convertida'])
    
    if variable_seleccionada and variable_seleccionada != "Todas las variables":
        # Filtrar columnas de la variable seleccionada
        estrategias_clave = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada][:6]
    else:
        # Estrategias clave para seguimiento temporal (todas las variables)
        estrategias_clave = [
            col for col in dummy_cols 
            if any(palabra in col.lower() for palabra in ['meme', 'logotipo', 'testimonio', 'plain_folks', 'orquestacion'])
        ][:6]
    
    if not estrategias_clave:
        return None, []
    
    # Agrupar por fecha
    df_agrupado = df_temporal.groupby('Fecha_convertida')[estrategias_clave].sum().reset_index()
    
    return df_agrupado, estrategias_clave

def analisis_propaganda_candidatos(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de técnicas de propaganda por candidato"""
    if variable_seleccionada and variable_seleccionada != "Todas las variables":
        propaganda_cols = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada]
    else:
        propaganda_cols = [col for col in dummy_cols if 'institute' in col.lower() or 'propaganda' in col.lower()]
    
    if not propaganda_cols or 'Candidato' not in df.columns:
        return None
    
    candidatos = df['Candidato'].unique()
    datos_propaganda = []
    
    for candidato in candidatos:
        df_cand = df[df['Candidato'] == candidato]
        total_posts = len(df_cand)
        
        for col in propaganda_cols:
            if col in df.columns:
                usos = df_cand[col].sum()
                porcentaje = (usos / total_posts) * 100 if total_posts > 0 else 0
                
                if '__' in col:
                    variable = col.split('__')[0].replace('_', ' ').title()
                    tecnica = col.split('__')[1].replace('_', ' ').title()
                    nombre_completo = f"{variable} - {tecnica}"
                else:
                    nombre_completo = col
                
                datos_propaganda.append({
                    'Candidato': candidato,
                    'Técnica': nombre_completo,
                    'Usos': usos,
                    'Porcentaje': porcentaje
                })
    
    return pd.DataFrame(datos_propaganda)

def analisis_plain_folks(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis detallado de la estrategia Plain-folks"""
    if variable_seleccionada and variable_seleccionada != "Todas las variables":
        plain_folks_cols = [col for col in dummy_cols if col.split('__')[0] == variable_seleccionada and 'plain' in col.lower()]
    else:
        plain_folks_cols = [col for col in dummy_cols if 'plain' in col.lower() or 'pueblo' in col.lower()]
    
    if not plain_folks_cols:
        return None
    
    # Filtrar posts que usan Plain-folks
    df_plain = df[df[plain_folks_cols].sum(axis=1) > 0].copy()
    
    resultados = {}
    
    # Por candidato
    if 'Candidato' in df.columns and len(df_plain) > 0:
        candidatos_stats = []
        for candidato in df['Candidato'].unique():
            df_cand = df[df['Candidato'] == candidato]
            df_plain_cand = df_plain[df_plain['Candidato'] == candidato]
            
            total_posts = len(df_cand)
            plain_posts = len(df_plain_cand)
            porcentaje = (plain_posts / total_posts) * 100 if total_posts > 0 else 0
            
            candidatos_stats.append({
                'Candidato': candidato,
                'Total_Posts': total_posts,
                'Plain_Folks_Posts': plain_posts,
                'Porcentaje': porcentaje
            })
        
        resultados['candidatos'] = pd.DataFrame(candidatos_stats)
    
    return resultados

def analisis_plain_folks_por_contexto(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis detallado de Plain-folks según contexto y campaña"""
    resultados = {}
    
    # Buscar columnas de Plain-folks
    plain_folks_cols = [col for col in dummy_cols if 'plain_folks' in col.lower() or 'plain_folk' in col.lower()]
    
    if not plain_folks_cols:
        return None
    
    # Análisis por contexto y candidato
    if 'Contexto_de_la_imagen' in df.columns and 'Candidato' in df.columns:
        contextos = df['Contexto_de_la_imagen'].dropna().unique()
        candidatos = df['Candidato'].dropna().unique()
        
        # Crear tabla de Plain-folks por contexto y candidato
        tabla_contexto = []
        
        for contexto in contextos:
            df_contexto = df[df['Contexto_de_la_imagen'] == contexto]
            
            # Calcular totales por candidato
            for candidato in candidatos:
                df_cand_contexto = df_contexto[df_contexto['Candidato'] == candidato]
                
                if len(df_cand_contexto) > 0:
                    # Contar publicaciones con Plain-folks
                    plain_folks_count = 0
                    for col in plain_folks_cols:
                        if col in df_cand_contexto.columns:
                            plain_folks_count += df_cand_contexto[col].sum()
                    
                    total_publicaciones = len(df_cand_contexto)
                    porcentaje = (plain_folks_count / total_publicaciones * 100) if total_publicaciones > 0 else 0
                    
                    tabla_contexto.append({
                        'Contexto': contexto,
                        'Campaña': candidato,
                        f'{candidato}: publicaciones con Plain-folks (%)': f"{plain_folks_count} ({porcentaje:.1f} %)",
                        'Total_publicaciones': total_publicaciones,
                        'Plain_folks_count': plain_folks_count
                    })
        
        if tabla_contexto:
            df_contexto_resultado = pd.DataFrame(tabla_contexto)
            
            # Reorganizar para mostrar como en la imagen
            tabla_pivot = []
            for contexto in contextos:
                fila = {'Contexto': contexto, 'Campaña': 'Votantes' if 'Vía pública' in contexto else 'Familiares'}
                
                for candidato in candidatos:
                    datos_contexto = df_contexto_resultado[
                        (df_contexto_resultado['Contexto'] == contexto)
                    ]
                    
                    for _, row in datos_contexto.iterrows():
                        if candidato in row[f'{candidato}: publicaciones con Plain-folks (%)']:
                            fila[f'{candidato}: publicaciones con Plain-folks (%)'] = row[f'{candidato}: publicaciones con Plain-folks (%)']
                
                tabla_pivot.append(fila)
            
            resultados['plain_folks_contexto'] = pd.DataFrame(tabla_pivot)
    
    return resultados

def analisis_distribucion_propaganda_ipa(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de distribución general de recursos de propaganda según IPA"""
    resultados = {}
    
    # Recursos de propaganda IPA
    recursos_ipa = [
        'name_calling', 'glittering_generalities', 'transfer', 
        'testimonial', 'plain_folks', 'card_stacking', 'bandwagon'
    ]
    
    if 'Candidato' in df.columns:
        candidatos = df['Candidato'].dropna().unique()
        
        tabla_recursos = []
        
        for i, recurso in enumerate(recursos_ipa, 1):
            fila = {'Recurso de propaganda (IPA)': f"{i}. {recurso.replace('_', '-').title()}"}
            
            for candidato in candidatos:
                df_candidato = df[df['Candidato'] == candidato]
                
                # Buscar columnas que contengan este recurso
                recurso_cols = [col for col in dummy_cols if recurso.lower() in col.lower()]
                
                total_recurso = 0
                for col in recurso_cols:
                    if col in df_candidato.columns:
                        total_recurso += df_candidato[col].sum()
                
                total_publicaciones = len(df_candidato)
                porcentaje = (total_recurso / total_publicaciones * 100) if total_publicaciones > 0 else 0
                
                fila[f'{candidato}: Nº de publicaciones'] = total_recurso
                fila[f'{candidato}: %'] = f"{porcentaje:.1f}"
            
            tabla_recursos.append(fila)
        
        resultados['distribucion_ipa'] = pd.DataFrame(tabla_recursos)
    
    return resultados

def analisis_cruce_reglas_contexto(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de cruce entre reglas de propaganda y contexto de imagen"""
    resultados = {}
    
    # Buscar variables de reglas dominantes y contexto
    if 'Contexto_de_la_imagen' in df.columns:
        contextos = df['Contexto_de_la_imagen'].dropna().unique()
        
        # Buscar columnas de reglas de propaganda
        reglas_cols = [col for col in dummy_cols if any(regla in col.lower() for regla in 
                      ['name_calling', 'glittering', 'transfer', 'testimonial', 'plain_folks', 'card_stacking', 'bandwagon'])]
        
        tabla_cruce = []
        
        for contexto in contextos:
            df_contexto = df[df['Contexto_de_la_imagen'] == contexto]
            
            # Encontrar la regla dominante en este contexto
            regla_dominante = None
            max_count = 0
            
            for col in reglas_cols:
                if col in df_contexto.columns:
                    count = df_contexto[col].sum()
                    if count > max_count:
                        max_count = count
                        regla_dominante = col.split('__')[-1] if '__' in col else col
            
            # Determinar campaña dominante
            campana_dominante = "Votantes"  # Por defecto
            candidato_dominante = "Ambos"
            
            if 'Candidato' in df_contexto.columns:
                candidatos_count = df_contexto['Candidato'].value_counts()
                if len(candidatos_count) > 0:
                    candidato_dominante = candidatos_count.index[0]
            
            # Mapear contexto a campaña
            if 'personal' in contexto.lower():
                campana_dominante = "Familiares"
            elif 'profesional' in contexto.lower():
                campana_dominante = "Ninguno"
            elif 'pública' in contexto.lower() or 'publica' in contexto.lower():
                campana_dominante = "Votantes"
            
            tabla_cruce.append({
                'Contexto de la imagen': contexto,
                'Regla dominante (Domenach)': regla_dominante.replace('_', ' ').title() if regla_dominante else "N/A",
                'Nº publicaciones': len(df_contexto),
                'Candidato dominante': candidato_dominante
            })
        
        resultados['cruce_reglas_contexto'] = pd.DataFrame(tabla_cruce)
    
    return resultados

def analisis_aparicion_lider(df, dummy_cols, variable_seleccionada=None, formato_apa=False):
    """Análisis de aparición del líder según contexto y campaña"""
    resultados = {}
    
    # Buscar columnas relacionadas con aparición del líder
    aparicion_cols = [col for col in dummy_cols if 'aparicion' in col.lower() or 'lider' in col.lower() or 'acompañado' in col.lower()]
    
    if 'Contexto_de_la_imagen' in df.columns and aparicion_cols:
        contextos = df['Contexto_de_la_imagen'].dropna().unique()
        
        tabla_aparicion = []
        
        for contexto in contextos:
            df_contexto = df[df['Contexto_de_la_imagen'] == contexto]
            
            # Determinar si va acompañado
            acompanado = "Ninguno"
            if 'personal' in contexto.lower():
                acompanado = "Familiares"
            elif 'pública' in contexto.lower() or 'publica' in contexto.lower():
                acompanado = "Votantes"
            
            # Calcular presencia del líder
            presencia_lider = 0
            total_publicaciones = len(df_contexto)
            
            for col in aparicion_cols:
                if col in df_contexto.columns:
                    presencia_lider += df_contexto[col].sum()
            
            porcentaje_presencia = (presencia_lider / total_publicaciones * 100) if total_publicaciones > 0 else 0
            
            tabla_aparicion.append({
                'Contexto': contexto,
                '¿Va acompañado?': acompanado,
                'Presencia del líder (%)': f"{porcentaje_presencia:.0f}%",
                'Total': total_publicaciones
            })
        
        resultados['aparicion_lider'] = pd.DataFrame(tabla_aparicion)
    
    return resultados

def generar_tabla_contingencia_avanzada(df, var1, var2, incluir_porcentajes=True):
    """Genera tabla de contingencia avanzada con múltiples estadísticos"""
    try:
        # Crear tabla de contingencia básica
        tabla = pd.crosstab(df[var1], df[var2], margins=True)
        
        resultados = {'tabla_frecuencias': tabla}
        
        if incluir_porcentajes:
            # Tabla de porcentajes por filas
            tabla_pct_filas = pd.crosstab(df[var1], df[var2], normalize='index') * 100
            tabla_pct_filas = tabla_pct_filas.round(1)
            
            # Tabla de porcentajes por columnas
            tabla_pct_cols = pd.crosstab(df[var1], df[var2], normalize='columns') * 100
            tabla_pct_cols = tabla_pct_cols.round(1)
            
            # Tabla de porcentajes totales
            tabla_pct_total = pd.crosstab(df[var1], df[var2], normalize='all') * 100
            tabla_pct_total = tabla_pct_total.round(1)
            
            resultados.update({
                'tabla_porcentajes_filas': tabla_pct_filas,
                'tabla_porcentajes_columnas': tabla_pct_cols,
                'tabla_porcentajes_total': tabla_pct_total
            })
        
        # Prueba de chi-cuadrado si es posible
        try:
            from scipy.stats import chi2_contingency
            tabla_sin_margenes = tabla.iloc[:-1, :-1]  # Quitar totales
            chi2, p_valor, gl, esperados = chi2_contingency(tabla_sin_margenes)
            
            resultados['estadisticos'] = {
                'chi2': chi2,
                'p_valor': p_valor,
                'grados_libertad': gl,
                'significativo': p_valor < 0.05
            }
        except:
            pass
        
        return resultados
    
    except Exception as e:
        st.error(f"Error al generar tabla de contingencia: {str(e)}")
        return None

def crear_visualizacion_avanzada(df, tipo_analisis, **kwargs):
    """Crea visualizaciones avanzadas según el tipo de análisis"""
    figs = {}
    
    if tipo_analisis == "distribucion_ipa":
        # Gráfico de barras apiladas para recursos IPA
        if 'datos' in kwargs and not kwargs['datos'].empty:
            datos = kwargs['datos']
            
            # Preparar datos para gráfico
            candidatos = [col for col in datos.columns if 'Nº de publicaciones' in col]
            
            fig = go.Figure()
            
            for candidato_col in candidatos:
                candidato = candidato_col.split(':')[0]
                valores = datos[candidato_col].values
                recursos = datos['Recurso de propaganda (IPA)'].values
                
                fig.add_trace(go.Bar(
                    name=candidato,
                    x=recursos,
                    y=valores,
                    text=valores,
                    textposition='auto'
                ))
            
            fig.update_layout(
                title="Distribución de Recursos de Propaganda (IPA) por Candidato",
                xaxis_title="Recursos de Propaganda",
                yaxis_title="Número de Publicaciones",
                barmode='group',
                height=600
            )
            
            figs['distribucion_ipa'] = fig
    
    elif tipo_analisis == "plain_folks_contexto":
        # Gráfico de barras para Plain-folks por contexto
        if 'datos' in kwargs and not kwargs['datos'].empty:
            datos = kwargs['datos']
            
            fig = go.Figure()
            
            # Extraer datos para visualización
            contextos = datos['Contexto'].unique()
            candidatos_cols = [col for col in datos.columns if 'publicaciones con Plain-folks' in col]
            
            for col in candidatos_cols:
                candidato = col.split(':')[0]
                # Extraer números de los porcentajes
                valores = []
                for _, row in datos.iterrows():
                    texto = str(row[col])
                    if '(' in texto and ')' in texto:
                        porcentaje = texto.split('(')[1].split('%')[0]
                        try:
                            valores.append(float(porcentaje))
                        except:
                            valores.append(0)
                    else:
                        valores.append(0)
                
                fig.add_trace(go.Bar(
                    name=candidato,
                    x=datos['Contexto'],
                    y=valores,
                    text=[f"{v:.1f}%" for v in valores],
                    textposition='auto'
                ))
            
            fig.update_layout(
                title="Uso del recurso Plain-folks según contexto y campaña",
                xaxis_title="Contexto",
                yaxis_title="Porcentaje de publicaciones (%)",
                barmode='group',
                height=500
            )
            
            figs['plain_folks_contexto'] = fig
    
    return figs

# =====================================================
# FUNCIONES PARA MODO CLARO/OSCURO Y EXPORTACIÓN
# =====================================================

def aplicar_tema(tema="claro"):
    """Aplica tema claro u oscuro para las visualizaciones"""
    if tema == "oscuro":
        plt.style.use('dark_background')
        plt.rcParams.update({
            'axes.facecolor': '#2e2e2e',
            'figure.facecolor': '#1e1e1e',
            'text.color': 'white',
            'axes.labelcolor': 'white',
            'xtick.color': 'white',
            'ytick.color': 'white',
            'axes.edgecolor': 'white',
            'grid.color': '#404040'
        })
        return True
    else:
        # Modo claro - APA estándar
        plt.style.use('default')
        plt.rcParams.update({
            'font.family': 'serif',
            'font.serif': ['Times New Roman', 'DejaVu Serif', 'serif'],
            'font.size': 12,
            'axes.titlesize': 14,
            'axes.labelsize': 12,
            'xtick.labelsize': 11,
            'ytick.labelsize': 11,
            'legend.fontsize': 11,
            'figure.titlesize': 16,
            'axes.spines.top': False,
            'axes.spines.right': False,
            'axes.spines.left': True,
            'axes.spines.bottom': True,
            'axes.linewidth': 1.0,
            'axes.edgecolor': 'black',
            'axes.facecolor': 'white',
            'figure.facecolor': 'white',
            'axes.grid': True,
            'grid.alpha': 0.3,
            'grid.linewidth': 0.5,
            'grid.color': 'gray',
            'figure.figsize': (10, 6),
            'figure.dpi': 300,
            'savefig.dpi': 300,
            'savefig.bbox': 'tight',
            'savefig.facecolor': 'white',
        })
        return False

def crear_tabla_apa_docx(doc, df, titulo="Tabla"):
    """Crea una tabla en formato APA para documento DOCX"""
    # Agregar título de tabla
    titulo_para = doc.add_paragraph()
    titulo_run = titulo_para.add_run(titulo)
    titulo_run.font.name = 'Times New Roman'
    titulo_run.font.size = Pt(12)
    titulo_run.bold = True
    titulo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Crear tabla
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'
    
    # Configurar encabezados
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = str(column)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Agregar datos
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Aplicar bordes APA
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(1.5)
    
    doc.add_paragraph()  # Espacio después de la tabla
    return table

def exportar_a_docx(dataframes_dict, graficos_dict=None, titulo_documento="Análisis de Campaña Electoral"):
    """Exporta tablas y gráficos a un documento DOCX en formato APA"""
    doc = Document()
    
    # Configurar documento
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    
    # Título principal
    titulo = doc.add_heading(titulo_documento, 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
    
    # Fecha
    fecha_para = doc.add_paragraph()
    fecha_run = fecha_para.add_run(f"Fecha de análisis: {datetime.now().strftime('%d de %B de %Y')}")
    fecha_run.font.name = 'Times New Roman'
    fecha_run.font.size = Pt(12)
    fecha_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espacio
    
    # Agregar tablas
    for i, (nombre, df) in enumerate(dataframes_dict.items(), 1):
        if not df.empty:
            titulo_tabla = f"Tabla {i}. {nombre}"
            crear_tabla_apa_docx(doc, df, titulo_tabla)
    
    # Guardar en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def obtener_css_tema(tema_oscuro=False):
    """Genera CSS personalizado según el tema seleccionado"""
    if tema_oscuro:
        return """
        <style>
            .stApp {
                background-color: #1e1e1e;
                color: white;
            }
            .main-header {
                font-size: 2.5rem;
                font-weight: bold;
                color: #4dabf7;
                text-align: center;
                margin-bottom: 2rem;
            }
            .section-header {
                font-size: 1.5rem;
                font-weight: bold;
                color: #74c0fc;
                margin-top: 2rem;
                margin-bottom: 1rem;
                border-bottom: 2px solid #364fc7;
                padding-bottom: 0.5rem;
            }
            .metric-card {
                background-color: #2c2c2c;
                padding: 1rem;
                border-radius: 10px;
                border-left: 4px solid #4dabf7;
                margin: 0.5rem 0;
            }
            .apa-table {
                font-family: 'Times New Roman', serif;
                font-size: 12px;
                background-color: #2c2c2c;
                color: white;
                border-collapse: collapse;
            }
            .apa-table th {
                border-top: 2px solid white;
                border-bottom: 1px solid white;
                padding: 8px;
                text-align: center;
                font-weight: bold;
                background-color: #404040;
            }
            .apa-table td {
                padding: 6px;
                text-align: center;
                border-bottom: none;
            }
            .apa-table .final-row td {
                border-bottom: 1px solid white;
            }
        </style>
        """
    else:
        return """
        <style>
            .main-header {
                font-size: 2.5rem;
                font-weight: bold;
                color: #1f77b4;
                text-align: center;
                margin-bottom: 2rem;
            }
            .section-header {
                font-size: 1.5rem;
                font-weight: bold;
                color: #2e86ab;
                margin-top: 2rem;
                margin-bottom: 1rem;
                border-bottom: 2px solid #e6f3ff;
                padding-bottom: 0.5rem;
            }
            .metric-card {
                background-color: #f8f9fa;
                padding: 1rem;
                border-radius: 10px;
                border-left: 4px solid #1f77b4;
                margin: 0.5rem 0;
            }
            .stDataFrame {
                border: 1px solid #e6e6e6;
                border-radius: 5px;
            }
            .apa-table {
                font-family: 'Times New Roman', serif;
                font-size: 12px;
                background-color: white;
                border-collapse: collapse;
            }
            .apa-table th {
                border-top: 2px solid black;
                border-bottom: 1px solid black;
                padding: 8px;
                text-align: center;
                font-weight: bold;
            }
            .apa-table td {
                padding: 6px;
                text-align: center;
                border-bottom: none;
            }
            .apa-table .final-row td {
                border-bottom: 1px solid black;
            }
            .download-button {
                background-color: #28a745;
                color: white;
                border: none;
                padding: 0.5rem 1rem;
                border-radius: 5px;
                cursor: pointer;
                margin: 0.5rem 0;
            }
        </style>
        """

# =====================================================
# FUNCIÓN PRINCIPAL
# =====================================================

def main():
    """Función principal de la aplicación Streamlit"""
    
    # =====================================================
    # CONFIGURACIÓN DEL TEMA DE LA PÁGINA
    # =====================================================
    
    # Aplicar CSS por defecto (será sobrescrito por el tema seleccionado)
    st.markdown("""
    <style>
        .main-header {
            font-size: 2.5rem;
            font-weight: bold;
            color: #1f77b4;
            text-align: center;
            margin-bottom: 2rem;
        }
        .section-header {
            font-size: 1.5rem;
            font-weight: bold;
            color: #2e86ab;
            margin-top: 2rem;
            margin-bottom: 1rem;
            border-bottom: 2px solid #e6f3ff;
            padding-bottom: 0.5rem;
        }
        .metric-card {
            background-color: #f8f9fa;
            padding: 1rem;
            border-radius: 10px;
            border-left: 4px solid #1f77b4;
            margin: 0.5rem 0;
        }
        .stDataFrame {
            border: 1px solid #e6e6e6;
            border-radius: 5px;
        }
        .apa-table {
            font-family: 'Times New Roman', serif;
            font-size: 12px;
            background-color: white;
            border-collapse: collapse;
        }
        .apa-table th {
            border-top: 2px solid black;
            border-bottom: 1px solid black;
            padding: 8px;
            text-align: center;
            font-weight: bold;
        }
        .apa-table td {
            padding: 6px;
            text-align: center;
            border-bottom: none;
        }
        .apa-table .final-row td {
            border-bottom: 1px solid black;
        }
        .download-button {
            background-color: #28a745;
            color: white;
            border: none;
            padding: 0.5rem 1rem;
            border-radius: 5px;
            cursor: pointer;
            margin: 0.5rem 0;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # =====================================================
    # CONFIGURACIÓN INICIAL
    # =====================================================
    
    st.markdown('<div class="main-header">📊 Análisis Avanzado de Campaña Electoral</div>', unsafe_allow_html=True)
    
    # Cargar datos
    df, dummy_cols = cargar_datos()
    
    if df is None or len(dummy_cols) == 0:
        st.error("❌ No se pudieron cargar los datos o no se encontraron columnas dummy.")
        st.info("Asegúrate de que el archivo 'recodificado.xlsx' esté en el directorio correcto.")
        return
      # =====================================================
    # CONFIGURACIÓN GLOBAL EN SIDEBAR
    # =====================================================
    
    st.sidebar.header("⚙️ Configuración General")
    
    # Toggle para formato APA
    formato_apa = st.sidebar.checkbox(
        "📋 Activar formato APA académico",
        value=False,
        help="Activa el formato de tablas según estándares APA para publicaciones académicas"
    )
    
    # Toggle para modo oscuro
    modo_oscuro = st.sidebar.checkbox(
        "🌙 Modo oscuro",
        value=False,
        help="Activa el modo oscuro para visualizaciones (no compatible con exportación APA)"
    )
    
    # Aplicar tema
    tema_aplicado = aplicar_tema("oscuro" if modo_oscuro else "claro")
    
    # Aplicar CSS según el tema
    st.markdown(obtener_css_tema(modo_oscuro), unsafe_allow_html=True)
    
    # Selección de variable principal
    variables_principales = obtener_variables_principales(dummy_cols)
    variable_seleccionada = st.sidebar.selectbox(
        "🔍 Seleccionar variable para análisis:",
        ["Todas las variables"] + variables_principales,
        help="Selecciona una variable específica para análisis univariado"
    )
    
    # Selección múltiple de categorías (si se seleccionó una variable específica)
    categorias_seleccionadas = []
    if variable_seleccionada != "Todas las variables":
        categorias_disponibles = obtener_categorias_de_variable(dummy_cols, variable_seleccionada)
        if categorias_disponibles:
            nombres_categorias = [cat[1].replace('_', ' ').title() for cat in categorias_disponibles]
            
            # Usar multiselect para selección múltiple
            categorias_display = st.sidebar.multiselect(
                f"📂 Categorías de {variable_seleccionada.replace('_', ' ').title()}:",
                options=nombres_categorias,
                default=[],
                help="Selecciona una o múltiples categorías para análisis. Si no seleccionas ninguna, se incluirán todas."
            )
            
            if categorias_display:
                # Convertir nombres display a nombres originales
                for cat_display in categorias_display:
                    for cat_orig, cat_orig_name in categorias_disponibles:
                        if cat_orig_name.replace('_', ' ').title() == cat_display:
                            categorias_seleccionadas.append(cat_orig_name)
                            break
            else:
                # Si no se selecciona nada, incluir todas las categorías
                categorias_seleccionadas = ["Todas las categorías"]
    
    # Filtros adicionales
    st.sidebar.header("🔧 Filtros Adicionales")
    
    candidatos_disponibles = ["Todos"] + list(df['Candidato'].unique()) if 'Candidato' in df.columns else ["Todos"]
    candidato_seleccionado = st.sidebar.selectbox("👤 Candidato:", candidatos_disponibles)
    
    # Filtro por fecha
    if 'Fecha_convertida' in df.columns:
        df_con_fecha = df.dropna(subset=['Fecha_convertida'])
        if len(df_con_fecha) > 0:
            fecha_min = df_con_fecha['Fecha_convertida'].min().date()
            fecha_max = df_con_fecha['Fecha_convertida'].max().date()
            
            rango_fechas = st.sidebar.date_input(
                "📅 Rango de fechas:",
                value=(fecha_min, fecha_max),
                min_value=fecha_min,
                max_value=fecha_max
            )
    
    # =====================================================
    # APLICAR FILTROS A LOS DATOS
    # =====================================================
    
    df_filtrado = df.copy()
    
    # Filtrar por candidato
    if candidato_seleccionado != "Todos":
        df_filtrado = df_filtrado[df_filtrado['Candidato'] == candidato_seleccionado]
    
    # Filtrar por fecha
    if 'Fecha_convertida' in df.columns and 'rango_fechas' in locals() and len(rango_fechas) == 2:
        fecha_inicio, fecha_fin = rango_fechas
        df_filtrado = df_filtrado[
            (df_filtrado['Fecha_convertida'].dt.date >= fecha_inicio) &
            (df_filtrado['Fecha_convertida'].dt.date <= fecha_fin)
        ]
      # Aplicar filtro de variable/categoría
    df_filtrado, dummy_cols_filtradas = filtrar_datos_por_seleccion(
        df_filtrado, dummy_cols, variable_seleccionada, categorias_seleccionadas
    )
    
    # Mostrar información de filtros aplicados
    st.sidebar.markdown("---")
    st.sidebar.markdown("📊 **Datos filtrados:**")
    st.sidebar.metric("Total de registros", len(df_filtrado))
    
    if variable_seleccionada != "Todas las variables":
        st.sidebar.info(f"🎯 **Análisis enfocado en:** {variable_seleccionada.replace('_', ' ').title()}")
        if categorias_seleccionadas and "Todas las categorías" not in categorias_seleccionadas:
            categorias_text = ", ".join([cat.replace('_', ' ').title() for cat in categorias_seleccionadas])
            st.sidebar.info(f"📂 **Categorías:** {categorias_text}")
    
    if formato_apa:
        st.sidebar.success("📋 Formato APA activo")
    
    # Verificar que hay datos después del filtrado
    if len(df_filtrado) == 0:
        st.warning("⚠️ No hay datos que coincidan con los filtros seleccionados.")
        return
      # =====================================================
    # SECCIÓN 1: RANKING DE CATEGORÍAS TOP
    # =====================================================
    
    st.markdown('<div class="section-header">🏆 Ranking de Categorías Más Utilizadas</div>', unsafe_allow_html=True)
    
    # Configuración del ranking
    col_config1, col_config2, col_config3 = st.columns([1, 1, 1])
    
    with col_config1:
        n_top = st.slider(
            "Número de categorías a mostrar:",
            min_value=5,
            max_value=20,
            value=10,
            step=1
        )
    
    with col_config2:
        mostrar_por_candidato = st.checkbox(
            "📊 Mostrar desglose por candidato",
            value=True,
            help="Muestra la distribución de cada categoría por candidato además del total"
        )
    
    with col_config3:
        tamaño_grafico = st.selectbox(
            "📏 Tamaño del gráfico:",
            ["Normal", "Grande", "Extra Grande"],
            index=0,
            help="Selecciona el tamaño del gráfico para mejor visualización"
        )
    
    # Determinar altura según selección
    altura_grafico = {
        "Normal": 500,
        "Grande": 700,
        "Extra Grande": 900
    }[tamaño_grafico]
    
    # Crear ranking
    if mostrar_por_candidato:
        df_top = crear_ranking_por_candidato_y_total(df_filtrado, dummy_cols_filtradas, variable_seleccionada, n_top, formato_apa)
    else:
        df_top = crear_ranking_por_variable(df_filtrado, dummy_cols_filtradas, variable_seleccionada, n_top, formato_apa)
    
    if len(df_top) > 0:
        # Mostrar tabla con formato seleccionado
        titulo_ranking = f"Top {n_top} Categorías"
        if variable_seleccionada != "Todas las variables":
            titulo_ranking += f" - {variable_seleccionada.replace('_', ' ').title()}"
        
        mostrar_tabla_con_formato(df_top, titulo_ranking, formato_apa)
        
        # Opción para expandir tabla
        with st.expander("📋 Ver tabla expandida"):
            st.dataframe(df_top, use_container_width=True, height=400)
        
        # Gráfico de barras con tamaño ajustable
        col_grafico1, col_grafico2 = st.columns([4, 1])
        
        with col_grafico1:
            fig = px.bar(
                df_top.head(10),
                x='Total' if mostrar_por_candidato else 'Frecuencia',
                y='Categoría',                orientation='h',
                title=titulo_ranking,
                labels={'Total': 'Número de Usos Total', 'Frecuencia': 'Número de Usos', 'Categoría': 'Categoría'},
                height=altura_grafico
            )
            
            fig.update_layout(
                yaxis={'categoryorder': 'total ascending'},
                xaxis_title="Número de Usos",
                yaxis_title="",
                title_x=0.5,
                margin=dict(l=200, r=50, t=80, b=50)  # Más margen izquierdo para etiquetas largas
            )
            # Mejorar legibilidad de etiquetas
            fig.update_yaxes(tickfont=dict(size=10))
            fig.update_xaxes(tickfont=dict(size=12))
            
            st.plotly_chart(fig, use_container_width=True)
        
        with col_grafico2:
            if st.button("🔍 Ver en pantalla completa", use_container_width=True):
                # Crear gráfico en pantalla completa
                fig_full = px.bar(
                    df_top,
                    x='Total' if mostrar_por_candidato else 'Frecuencia',
                    y='Categoría',
                    orientation='h',
                    title=f"{titulo_ranking} - Vista Expandida",
                    labels={'Total': 'Número de Usos Total', 'Frecuencia': 'Número de Usos', 'Categoría': 'Categoría'},
                    height=max(800, len(df_top) * 40)  # Altura dinámica según número de categorías
                )
                fig_full.update_layout(
                    yaxis={'categoryorder': 'total ascending'},
                    xaxis_title="Número de Usos",
                    yaxis_title="",
                    title_x=0.5,
                    margin=dict(l=300, r=50, t=100, b=50),  # Márgenes amplios
                    font=dict(size=14)  # Fuente más grande
                )
                
                st.plotly_chart(fig_full, use_container_width=True)
        
        # Si se muestra por candidato, agregar gráfico de barras apiladas
        if mostrar_por_candidato and 'Candidato' in df_filtrado.columns:
            st.markdown("### 📊 Distribución por Candidato")
            
            # Preparar datos para gráfico apilado
            candidatos = [col for col in df_top.columns if col not in ['Categoría', 'Total', 'Porcentaje']]
            
            if len(candidatos) > 0:
                fig_stack = go.Figure()
                
                for candidato in candidatos:
                    fig_stack.add_trace(go.Bar(
                        name=candidato,
                        x=df_top['Categoría'][:10],  # Solo top 10 para legibilidad
                        y=df_top[candidato][:10],
                        text=df_top[candidato][:10],
                        textposition='inside'
                    ))
                
                fig_stack.update_layout(
                    barmode='stack',
                    title="Distribución de Categorías por Candidato",
                    xaxis_title="Categorías",
                    yaxis_title="Número de Usos",
                    height=altura_grafico,
                    xaxis_tickangle=-45,
                    margin=dict(l=50, r=50, t=80, b=150)  # Margen inferior para etiquetas rotadas
                )
                
                st.plotly_chart(fig_stack, use_container_width=True)
    else:
        st.info("No hay datos para mostrar en el ranking.")
      # =====================================================
    # SECCIÓN 2: EVOLUCIÓN TEMPORAL
    # =====================================================
    
    st.markdown('<div class="section-header">📈 Evolución Temporal de Estrategias</div>', unsafe_allow_html=True)
    
    # Configuración para gráficos temporales
    col_temp1, col_temp2 = st.columns([3, 1])
    
    with col_temp2:
        tamaño_temporal = st.selectbox(
            "📏 Tamaño del gráfico temporal:",
            ["Normal", "Grande", "Extra Grande"],
            index=0,
            key="temporal_size"
        )
        
        altura_temporal = {
            "Normal": 500,
            "Grande": 700,
            "Extra Grande": 900
        }[tamaño_temporal]
    
    resultado_temporal = analisis_evolucion_temporal(df_filtrado, dummy_cols_filtradas, variable_seleccionada, formato_apa)
    
    if resultado_temporal[0] is not None:
        df_temporal, estrategias_clave = resultado_temporal
        
        if len(df_temporal) > 0 and len(estrategias_clave) > 0:
            with col_temp1:
                # Gráfico de líneas interactivo
                fig = go.Figure()
                
                for estrategia in estrategias_clave:
                    if estrategia in df_temporal.columns:
                        nombre_limpio = estrategia.split('__')[1].replace('_', ' ').title() if '__' in estrategia else estrategia
                        fig.add_trace(go.Scatter(
                            x=df_temporal['Fecha_convertida'],
                            y=df_temporal[estrategia],
                            mode='lines+markers',
                            name=nombre_limpio,
                            line=dict(width=3),
                            marker=dict(size=8)
                        ))
                
                titulo_temporal = "Evolución Temporal de Estrategias"
                if variable_seleccionada != "Todas las variables":
                    titulo_temporal += f" - {variable_seleccionada.replace('_', ' ').title()}"
                
                fig.update_layout(
                    title=titulo_temporal,
                    xaxis_title="Fecha",
                    yaxis_title="Número de Usos",
                    height=altura_temporal,
                    hovermode='x unified',
                    title_x=0.5,
                    margin=dict(l=50, r=50, t=80, b=50),
                    legend=dict(
                        orientation="h",
                        yanchor="bottom",
                        y=1.02,
                        xanchor="right",
                        x=1
                    )
                )
                
                st.plotly_chart(fig, use_container_width=True)
            
            # Botón para vista expandida
            if st.button("🔍 Ver evolución temporal en pantalla completa", key="temp_full"):
                fig_full_temp = go.Figure()
                
                for estrategia in estrategias_clave:
                    if estrategia in df_temporal.columns:
                        nombre_limpio = estrategia.split('__')[1].replace('_', ' ').title() if '__' in estrategia else estrategia
                        fig_full_temp.add_trace(go.Scatter(
                            x=df_temporal['Fecha_convertida'],
                            y=df_temporal[estrategia],
                            mode='lines+markers',
                            name=nombre_limpio,
                            line=dict(width=4),
                            marker=dict(size=10)
                        ))
                
                fig_full_temp.update_layout(
                    title=f"{titulo_temporal} - Vista Expandida",
                    xaxis_title="Fecha",
                    yaxis_title="Número de Usos",
                    height=800,
                    hovermode='x unified',
                    title_x=0.5,
                    font=dict(size=14),
                    margin=dict(l=80, r=80, t=100, b=80)                )
                
                st.plotly_chart(fig_full_temp, use_container_width=True)
            
            # Estadísticas temporales
            col1, col2 = st.columns(2)
            
            with col1:
                estadisticas = []
                for estrategia in estrategias_clave:
                    if estrategia in df_temporal.columns:
                        valores = df_temporal[estrategia]
                        nombre_limpio = estrategia.split('__')[1].replace('_', ' ').title() if '__' in estrategia else estrategia
                        
                        estadisticas.append({
                            'Estrategia': nombre_limpio,
                            'Total': int(valores.sum()),
                            'Promedio': round(valores.mean(), 2),
                            'Máximo': int(valores.max()),
                            'Mínimo': int(valores.min())
                        })
                
                df_stats = pd.DataFrame(estadisticas)
                mostrar_tabla_con_formato(df_stats, "📊 Estadísticas Descriptivas", formato_apa)
            
            with col2:
                st.subheader("📈 Tendencias")
                for estrategia in estrategias_clave:
                    if estrategia in df_temporal.columns:
                        nombre_limpio = estrategia.split('__')[1].replace('_', ' ').title() if '__' in estrategia else estrategia
                        total_usos = df_temporal[estrategia].sum()
                        tendencia = "📈" if df_temporal[estrategia].iloc[-1] > df_temporal[estrategia].iloc[0] else "📉"
                        
                        st.metric(
                            label=nombre_limpio,
                            value=int(total_usos),
                            delta=f"{tendencia} Tendencia"
                        )
        else:
            st.info("No hay datos temporales suficientes para el análisis.")
    else:
        st.info("No se encontraron estrategias clave para el análisis temporal.")
    
    # =====================================================
    # SECCIÓN 3: TABLA CRUZADA ENTRE VARIABLES
    # =====================================================
    
    st.markdown('<div class="section-header">🔄 Análisis de Cruces entre Variables</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Seleccionar Variables para Cruce")
        
        # Filtrar variables con suficientes datos
        variables_disponibles = [col for col in dummy_cols_filtradas if df_filtrado[col].sum() >= 3]
        
        if len(variables_disponibles) >= 2:
            var1 = st.selectbox(
                "Variable 1:",
                variables_disponibles,
                format_func=lambda x: f"{x.split('__')[0].replace('_', ' ').title()} - {x.split('__')[1].replace('_', ' ').title()}" if '__' in x else x
            )
            
            var2 = st.selectbox(
                "Variable 2:",
                [v for v in variables_disponibles if v != var1],
                format_func=lambda x: f"{x.split('__')[0].replace('_', ' ').title()} - {x.split('__')[1].replace('_', ' ').title()}" if '__' in x else x
            )
            
            tipo_visualizacion = st.radio(
                "Tipo de visualización:",
                ["Heatmap", "Tabla interactiva", "Ambos"]
            )
        else:
            st.warning("No hay suficientes variables con datos para realizar cruces.")
            var1, var2 = None, None
    
    with col2:
        if len(variables_disponibles) >= 2 and var1 and var2:
            tabla, tabla_pct, chi2_stat, p_value = crear_tabla_cruzada(df_filtrado, var1, var2, formato_apa)
            
            if tabla is not None:
                st.subheader("Estadísticas del Cruce")
                
                # Mostrar estadísticas
                total_casos = tabla.iloc[:-1, :-1].sum().sum()
                st.metric("Total de casos analizados", int(total_casos))
                
                if chi2_stat is not None:
                    col_chi1, col_chi2 = st.columns(2)
                    with col_chi1:
                        st.metric("Chi-cuadrado", f"{chi2_stat:.3f}")
                    with col_chi2:
                        st.metric("Valor p", f"{p_value:.3f}")
                
                # Mostrar tablas
                if tipo_visualizacion in ["Tabla interactiva", "Ambos"]:
                    mostrar_tabla_con_formato(tabla, "Tabla de Contingencia", formato_apa)
                    mostrar_tabla_con_formato(tabla_pct, "Tabla de Porcentajes", formato_apa)
                  # Mostrar heatmap
                if tipo_visualizacion in ["Heatmap", "Ambos"]:
                    # Configuración del heatmap
                    col_heat1, col_heat2 = st.columns([3, 1])
                    
                    with col_heat2:
                        tamaño_heatmap = st.selectbox(
                            "📏 Tamaño del heatmap:",
                            ["Normal", "Grande", "Extra Grande"],
                            index=0,
                            key="heatmap_size"
                        )
                        
                        altura_heatmap = {
                            "Normal": 500,
                            "Grande": 700,
                            "Extra Grande": 900
                        }[tamaño_heatmap]
                    
                    with col_heat1:
                        fig = px.imshow(
                            tabla_pct.iloc[:-1, :-1],  # Excluir márgenes
                            labels=dict(x="Variable 2", y="Variable 1", color="Porcentaje"),
                            title="Heatmap de Porcentajes",
                            height=altura_heatmap
                        )
                        fig.update_layout(
                            title_x=0.5,
                            margin=dict(l=150, r=50, t=80, b=50)
                        )
                        st.plotly_chart(fig, use_container_width=True)
                    
                    # Botón para vista expandida del heatmap
                    if st.button("🔍 Ver heatmap en pantalla completa", key="heat_full"):
                        fig_heat_full = px.imshow(
                            tabla_pct.iloc[:-1, :-1],
                            labels=dict(x="Variable 2", y="Variable 1", color="Porcentaje"),
                            title="Heatmap de Porcentajes - Vista Expandida",
                            height=800
                        )
                        fig_heat_full.update_layout(
                            title_x=0.5,
                            font=dict(size=14),
                            margin=dict(l=200, r=100, t=100, b=100)
                        )
                        st.plotly_chart(fig_heat_full, use_container_width=True)
      # =====================================================
    # SECCIÓN 4: ANÁLISIS DE PROPAGANDA
    # =====================================================
    
    st.markdown('<div class="section-header">📢 Análisis de Técnicas de Propaganda</div>', unsafe_allow_html=True)
    
    # Configuración para análisis de propaganda
    col_prop_config = st.columns([3, 1])
    
    with col_prop_config[1]:
        tamaño_propaganda = st.selectbox(
            "📏 Tamaño del gráfico:",
            ["Normal", "Grande", "Extra Grande"],
            index=0,
            key="prop_size"
        )
        
        altura_propaganda = {
            "Normal": 500,
            "Grande": 700,
            "Extra Grande": 900
        }[tamaño_propaganda]
    
    datos_propaganda = analisis_propaganda_candidatos(df_filtrado, dummy_cols_filtradas, variable_seleccionada, formato_apa)
    
    if datos_propaganda is not None and len(datos_propaganda) > 0:
        col1, col2 = st.columns(2)
        
        with col1:
            # Gráfico por candidato
            fig = px.bar(
                datos_propaganda,
                x='Candidato',
                y='Porcentaje',
                color='Técnica',
                title="Uso de Técnicas por Candidato (%)",
                labels={'Porcentaje': 'Porcentaje de Posts (%)'},
                height=altura_propaganda
            )
            fig.update_layout(
                title_x=0.5,
                margin=dict(l=50, r=50, t=80, b=100),
                xaxis_tickangle=-45
            )
            st.plotly_chart(fig, use_container_width=True)
            
            # Botón para vista expandida
            if st.button("🔍 Ver análisis de propaganda en pantalla completa", key="prop_full"):
                fig_prop_full = px.bar(
                    datos_propaganda,
                    x='Candidato',
                    y='Porcentaje',
                    color='Técnica',
                    title="Uso de Técnicas por Candidato (%) - Vista Expandida",
                    labels={'Porcentaje': 'Porcentaje de Posts (%)'},
                    height=800
                )
                fig_prop_full.update_layout(
                    title_x=0.5,
                    font=dict(size=14),
                    margin=dict(l=80, r=80, t=100, b=120),
                    xaxis_tickangle=-45
                )
                st.plotly_chart(fig_prop_full, use_container_width=True)
        
        with col2:
            titulo_propaganda = "Técnicas de Propaganda por Candidato"
            if variable_seleccionada != "Todas las variables":
                titulo_propaganda += f" - {variable_seleccionada.replace('_', ' ').title()}"
            
            mostrar_tabla_con_formato(datos_propaganda.round(2), titulo_propaganda, formato_apa)
            
            # Tabla expandida
            with st.expander("📋 Ver tabla expandida"):
                st.dataframe(datos_propaganda.round(2), use_container_width=True, height=300)
    else:
        st.info("No se encontraron datos de técnicas de propaganda.")
    
    # =====================================================
    # SECCIÓN 5: ANÁLISIS PLAIN-FOLKS
    # =====================================================
    
    st.markdown('<div class="section-header">👥 Análisis Detallado: Estrategia Plain-Folks</div>', unsafe_allow_html=True)
    
    resultados_plain = analisis_plain_folks(df_filtrado, dummy_cols_filtradas, variable_seleccionada, formato_apa)
    
    if resultados_plain:
        tab1, tab2 = st.tabs(["Por Candidato", "Detalles Contextuales"])
        
        with tab1:
            if 'candidatos' in resultados_plain:
                df_plain_cand = resultados_plain['candidatos']
                
                col1, col2 = st.columns(2)
                
                with col1:
                    fig = px.bar(
                        df_plain_cand,
                        x='Candidato',
                        y='Porcentaje',
                        title="Uso de Estrategia Plain-Folks por Candidato",                        labels={'Porcentaje': 'Porcentaje de Posts (%)'}
                    )
                    st.plotly_chart(fig, use_container_width=True)
                
                with col2:
                    titulo_plain = "Plain-Folks por Candidato"
                    if variable_seleccionada != "Todas las variables":
                        titulo_plain += f" - {variable_seleccionada.replace('_', ' ').title()}"
                    
                    mostrar_tabla_con_formato(df_plain_cand.round(2), titulo_plain, formato_apa)
        
        with tab2:
            contextos_encontrados = [k for k in resultados_plain.keys() if k.startswith('contexto_')]
            
            if contextos_encontrados:
                for contexto_key in contextos_encontrados[:3]:  # Máximo 3 contextos
                    contexto_nombre = contexto_key.replace('contexto_', '')
                    titulo_contexto = f"Plain-Folks en {contexto_nombre}"
                    
                    tabla_contexto = resultados_plain[contexto_key]
                    mostrar_tabla_con_formato(tabla_contexto, titulo_contexto, formato_apa)
            else:
                st.info("No se encontraron datos de contexto para Plain-Folks.")
    else:
        st.info("No se encontraron datos de la estrategia Plain-Folks.")
    
    # =====================================================
    # SECCIÓN 6: ANÁLISIS PLAIN-FOLKS POR CONTEXTO Y CAMPAÑA
    # =====================================================
    
    st.markdown('<div class="section-header">🍀 Análisis Plain-folks por Contexto y Campaña</div>', unsafe_allow_html=True)
    
    with st.expander("ℹ️ ¿Qué muestra este análisis?", expanded=False):
        st.write("""
        **Análisis detallado del uso del recurso Plain-folks según el contexto de la imagen y la campaña:**
        - Muestra el porcentaje de publicaciones que utilizan Plain-folks en cada contexto
        - Compara el uso entre diferentes candidatos
        - Analiza la efectividad del recurso según el escenario
        """)
    
    resultados_plain_contexto = analisis_plain_folks_por_contexto(df_filtrado, dummy_cols_filtradas, variable_seleccionada, formato_apa)
    
    if resultados_plain_contexto and 'plain_folks_contexto' in resultados_plain_contexto:
        df_plain_contexto = resultados_plain_contexto['plain_folks_contexto']
        
        if not df_plain_contexto.empty:
            # Opciones de visualización
            col_size6, col_download6 = st.columns([3, 1])
            
            with col_size6:
                tamano_tabla6 = st.selectbox(
                    "📏 Tamaño de tabla:",
                    ["Normal", "Compacta", "Expandida"],
                    key="tamano_plain_contexto"
                )
            
            # Mostrar tabla
            st.markdown("### 📊 Uso del recurso Plain-folks según contexto y campaña")
            
            if formato_apa:
                tabla_apa6 = aplicar_formato_apa_dataframe(df_plain_contexto, 
                    "Tabla: Uso del recurso Plain-folks según contexto y campaña")
                st.markdown(tabla_apa6, unsafe_allow_html=True)
            else:
                if tamano_tabla6 == "Compacta":
                    st.dataframe(df_plain_contexto, use_container_width=True, height=300)
                elif tamano_tabla6 == "Expandida":
                    st.dataframe(df_plain_contexto, use_container_width=True, height=600)
                else:
                    st.dataframe(df_plain_contexto, use_container_width=True)
            
            # Visualización
            figs_plain_contexto = crear_visualizacion_avanzada(df_filtrado, "plain_folks_contexto", datos=df_plain_contexto)
            
            if 'plain_folks_contexto' in figs_plain_contexto:
                col_graf6, col_config6 = st.columns([4, 1])
                
                with col_config6:
                    altura_graf6 = st.selectbox(
                        "📐 Altura del gráfico:",
                        [400, 500, 600, 700, 800],
                        index=1,
                        key="altura_plain_contexto"
                    )
                    
                    pantalla_completa6 = st.checkbox(
                        "🖥️ Modo pantalla completa",
                        key="fullscreen_plain_contexto"
                    )
                
                with col_graf6:
                    fig_plain_contexto = figs_plain_contexto['plain_folks_contexto']
                    fig_plain_contexto.update_layout(height=altura_graf6)
                    
                    if pantalla_completa6:
                        st.plotly_chart(fig_plain_contexto, use_container_width=True, height=altura_graf6, config={
                            'displayModeBar': True,
                            'displaylogo': False,
                            'modeBarButtonsToAdd': ['pan2d', 'select2d', 'lasso2d', 'autoScale2d', 'resetScale2d']
                        })
                    else:
                        st.plotly_chart(fig_plain_contexto, use_container_width=True)
        else:
            st.info("No se encontraron datos para el análisis de Plain-folks por contexto.")
    else:
        st.info("No se encontraron datos suficientes para este análisis.")
    
    # =====================================================
    # SECCIÓN 7: DISTRIBUCIÓN GENERAL DE RECURSOS IPA
    # =====================================================
    
    st.markdown('<div class="section-header">📊 Distribución General de Recursos de Propaganda (IPA)</div>', unsafe_allow_html=True)
    
    with st.expander("ℹ️ ¿Qué muestra este análisis?", expanded=False):
        st.write("""
        **Análisis de la distribución de los 7 recursos principales de propaganda según el IPA:**
        - Name-calling, Glittering generalities, Transfer, Testimonial, Plain folks, Card-stacking, Bandwagon
        - Compara el uso de cada recurso entre candidatos
        - Ideal para comparar el estilo comunicativo general
        """)
    
    resultados_ipa = analisis_distribucion_propaganda_ipa(df_filtrado, dummy_cols_filtradas, variable_seleccionada, formato_apa)
    
    if resultados_ipa and 'distribucion_ipa' in resultados_ipa:
        df_ipa = resultados_ipa['distribucion_ipa']
        
        if not df_ipa.empty:
            # Opciones de visualización
            col_size7, col_download7 = st.columns([3, 1])
            
            with col_size7:
                tamano_tabla7 = st.selectbox(
                    "📏 Tamaño de tabla:",
                    ["Normal", "Compacta", "Expandida"],
                    key="tamano_ipa"
                )
            
            # Mostrar tabla
            st.markdown("### 📊 Distribución general de recursos de propaganda (IPA)")
            
            if formato_apa:
                tabla_apa7 = aplicar_formato_apa_dataframe(df_ipa, 
                    "Tabla: Distribución general de recursos de propaganda (IPA)")
                st.markdown(tabla_apa7, unsafe_allow_html=True)
            else:
                if tamano_tabla7 == "Compacta":
                    st.dataframe(df_ipa, use_container_width=True, height=300)
                elif tamano_tabla7 == "Expandida":
                    st.dataframe(df_ipa, use_container_width=True, height=600)
                else:
                    st.dataframe(df_ipa, use_container_width=True)
            
            # Visualización
            figs_ipa = crear_visualizacion_avanzada(df_filtrado, "distribucion_ipa", datos=df_ipa)
            
            if 'distribucion_ipa' in figs_ipa:
                col_graf7, col_config7 = st.columns([4, 1])
                
                with col_config7:
                    altura_graf7 = st.selectbox(
                        "📐 Altura del gráfico:",
                        [500, 600, 700, 800, 900],
                        index=1,
                        key="altura_ipa"
                    )
                    
                    pantalla_completa7 = st.checkbox(
                        "🖥️ Modo pantalla completa",
                        key="fullscreen_ipa"
                    )
                
                with col_graf7:
                    fig_ipa = figs_ipa['distribucion_ipa']
                    fig_ipa.update_layout(height=altura_graf7)
                    
                    if pantalla_completa7:
                        st.plotly_chart(fig_ipa, use_container_width=True, height=altura_graf7, config={
                            'displayModeBar': True,
                            'displaylogo': False,
                            'modeBarButtonsToAdd': ['pan2d', 'select2d', 'lasso2d', 'autoScale2d', 'resetScale2d']
                        })
                    else:
                        st.plotly_chart(fig_ipa, use_container_width=True)
        else:
            st.info("No se encontraron datos para el análisis de distribución IPA.")
    else:
        st.info("No se encontraron datos suficientes para este análisis.")
    
    # =====================================================
    # SECCIÓN 8: CRUCE ENTRE REGLAS DE PROPAGANDA Y CONTEXTO
    # =====================================================
    
    st.markdown('<div class="section-header">✅ Cruce entre Reglas de Propaganda y Contexto de la Imagen</div>', unsafe_allow_html=True)
    
    with st.expander("ℹ️ ¿Por qué es importante este análisis?", expanded=False):
        st.write("""
        **Análisis que revela la efectividad contextual de las reglas de propaganda:**
        - Identifica qué reglas se activan en entornos más personales, profesionales o públicos
        - Proporciona una lectura del tono propagandístico según el marco escénico
        - Ayuda a entender la estrategia comunicativa contextual
        """)
    
    resultados_cruce = analisis_cruce_reglas_contexto(df_filtrado, dummy_cols_filtradas, variable_seleccionada, formato_apa)
    
    if resultados_cruce and 'cruce_reglas_contexto' in resultados_cruce:
        df_cruce = resultados_cruce['cruce_reglas_contexto']
        
        if not df_cruce.empty:
            # Opciones de visualización
            col_size8, col_download8 = st.columns([3, 1])
            
            with col_size8:
                tamano_tabla8 = st.selectbox(
                    "📏 Tamaño de tabla:",
                    ["Normal", "Compacta", "Expandida"],
                    key="tamano_cruce"
                )
            
            # Mostrar tabla
            st.markdown("### ✅ Cruce entre reglas de propaganda y contexto de la imagen")
            
            if formato_apa:
                tabla_apa8 = aplicar_formato_apa_dataframe(df_cruce, 
                    "Tabla: Cruce entre reglas de propaganda y contexto de la imagen")
                st.markdown(tabla_apa8, unsafe_allow_html=True)
            else:
                if tamano_tabla8 == "Compacta":
                    st.dataframe(df_cruce, use_container_width=True, height=300)
                elif tamano_tabla8 == "Expandida":
                    st.dataframe(df_cruce, use_container_width=True, height=600)
                else:
                    st.dataframe(df_cruce, use_container_width=True)
        else:
            st.info("No se encontraron datos para el análisis de cruce reglas-contexto.")
    else:
        st.info("No se encontraron datos suficientes para este análisis.")
    
    # =====================================================
    # SECCIÓN 9: APARICIÓN DEL LÍDER SEGÚN CONTEXTO
    # =====================================================
    
    st.markdown('<div class="section-header">📊 Aparición del Líder según Contexto y Campaña</div>', unsafe_allow_html=True)
    
    with st.expander("ℹ️ ¿Qué analiza esta tabla?", expanded=False):
        st.write("""
        **Variables cruzadas:**
        - Aparición del líder
        - Contexto de la imagen  
        - Aparición de terceras personas
        
        **Análisis de presencia del líder en diferentes contextos y su acompañamiento.**
        """)
    
    resultados_lider = analisis_aparicion_lider(df_filtrado, dummy_cols_filtradas, variable_seleccionada, formato_apa)
    
    if resultados_lider and 'aparicion_lider' in resultados_lider:
        df_lider = resultados_lider['aparicion_lider']
        
        if not df_lider.empty:
            # Opciones de visualización
            col_size9, col_download9 = st.columns([3, 1])
            
            with col_size9:
                tamano_tabla9 = st.selectbox(
                    "📏 Tamaño de tabla:",
                    ["Normal", "Compacta", "Expandida"],
                    key="tamano_lider"
                )
            
            # Mostrar tabla
            st.markdown("### 📊 Aparición del líder según contexto y campaña")
            
            if formato_apa:
                tabla_apa9 = aplicar_formato_apa_dataframe(df_lider, 
                    "Tabla: Aparición del líder según contexto y campaña")
                st.markdown(tabla_apa9, unsafe_allow_html=True)
            else:
                if tamano_tabla9 == "Compacta":
                    st.dataframe(df_lider, use_container_width=True, height=300)
                elif tamano_tabla9 == "Expandida":
                    st.dataframe(df_lider, use_container_width=True, height=600)
                else:
                    st.dataframe(df_lider, use_container_width=True)
        else:
            st.info("No se encontraron datos para el análisis de aparición del líder.")
    else:
        st.info("No se encontraron datos suficientes para este análisis.")
    
    # =====================================================
    # SECCIÓN 10: EXPORTACIÓN DE DATOS
    # =====================================================
    
    st.markdown('<div class="section-header">💾 Exportar Resultados</div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📥 Exportar Top Categorías", use_container_width=True):
            if 'df_top' in locals() and len(df_top) > 0:
                excel_data = exportar_a_excel(
                    {"Top_Categorias": df_top},
                    "top_categorias_campana"
                )
                
                st.download_button(
                    label="Descargar Excel - Top Categorías",
                    data=excel_data,
                    file_name="top_categorias_campana.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    
    with col2:
        if st.button("📥 Exportar Tabla Cruzada", use_container_width=True):
            if 'tabla' in locals() and tabla is not None:
                excel_data = exportar_a_excel(
                    {
                        "Tabla_Contingencia": tabla,
                        "Tabla_Porcentajes": tabla_pct
                    },
                    "tabla_cruzada_campana"
                )
                
                st.download_button(
                    label="Descargar Excel - Tabla Cruzada",
                    data=excel_data,
                    file_name="tabla_cruzada_campana.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    
    with col3:
        if st.button("📥 Exportar Análisis Propaganda", use_container_width=True):
            if 'datos_propaganda' in locals() and datos_propaganda is not None:
                excel_data = exportar_a_excel(
                    {"Propaganda_Candidatos": datos_propaganda},
                    "propaganda_campana"
                )
                
                st.download_button(
                    label="Descargar Excel - Propaganda",
                    data=excel_data,
                    file_name="propaganda_campana.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
    
    # Botón para exportar todo
    st.markdown("---")
    if st.button("📦 Exportar Análisis Completo", use_container_width=True):
        dataframes_completo = {}
        
        if 'df_top' in locals() and len(df_top) > 0:
            dataframes_completo["Top_Categorias"] = df_top
        
        if 'tabla' in locals() and tabla is not None:
            dataframes_completo["Tabla_Contingencia"] = tabla
            dataframes_completo["Tabla_Porcentajes"] = tabla_pct
        
        if 'datos_propaganda' in locals() and datos_propaganda is not None:
            dataframes_completo["Propaganda"] = datos_propaganda
        
        if 'df_temporal' in locals() and len(df_temporal) > 0:
            dataframes_completo["Evolucion_Temporal"] = df_temporal
        
        if dataframes_completo:
            excel_data = exportar_a_excel(dataframes_completo, "analisis_completo_campana")
            
            st.download_button(
                label="Descargar Excel - Análisis Completo",
                data=excel_data,
                file_name="analisis_completo_campana.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
        else:
            st.warning("No hay datos para exportar.")
    
    # Sección para exportación DOCX en formato APA
    st.markdown("---")
    st.markdown("### 📄 Exportación en Formato APA (DOCX)")
    
    col_docx1, col_docx2 = st.columns(2)
    
    with col_docx1:
        if st.button("📄 Exportar Tablas APA (DOCX)", use_container_width=True):
            # Solo exportar si el modo claro está activado (para mantener estándares APA)
            if not modo_oscuro:
                dataframes_apa = {}
                
                # Recopilar todas las tablas disponibles
                if 'df_top' in locals() and len(df_top) > 0:
                    dataframes_apa["Top Categorías de Estrategias"] = df_top
                
                if 'tabla' in locals() and tabla is not None:
                    dataframes_apa["Tabla de Contingencia"] = tabla
                    dataframes_apa["Tabla de Porcentajes"] = tabla_pct
                
                if 'datos_propaganda' in locals() and datos_propaganda is not None:
                    dataframes_apa["Análisis por Candidato"] = datos_propaganda
                
                if 'df_temporal' in locals() and len(df_temporal) > 0:
                    dataframes_apa["Evolución Temporal"] = df_temporal
                
                if 'df_plain_cand' in locals() and len(df_plain_cand) > 0:
                    dataframes_apa["Análisis Plain Folks por Candidato"] = df_plain_cand
                
                # Nuevas tablas agregadas
                if 'df_plain_contexto' in locals() and len(df_plain_contexto) > 0:
                    dataframes_apa["Plain-folks por Contexto y Campaña"] = df_plain_contexto
                
                if 'df_ipa' in locals() and len(df_ipa) > 0:
                    dataframes_apa["Distribución de Recursos IPA"] = df_ipa
                
                if 'df_cruce' in locals() and len(df_cruce) > 0:
                    dataframes_apa["Cruce Reglas-Contexto"] = df_cruce
                
                if 'df_lider' in locals() and len(df_lider) > 0:
                    dataframes_apa["Aparición del Líder por Contexto"] = df_lider
                
                if dataframes_apa:
                    # Generar documento DOCX
                    titulo_doc = f"Análisis de Campaña Electoral - {variable_seleccionada.replace('_', ' ').title() if variable_seleccionada != 'Todas las variables' else 'Análisis Completo'}"
                    docx_buffer = exportar_a_docx(dataframes_apa, titulo_documento=titulo_doc)
                    
                    st.download_button(
                        label="📄 Descargar Documento APA (DOCX)",
                        data=docx_buffer,
                        file_name=f"analisis_campana_apa_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.warning("No hay tablas disponibles para exportar.")
            else:
                st.warning("⚠️ La exportación APA requiere modo claro. Desactiva el modo oscuro para exportar en formato académico.")
    
    with col_docx2:
        st.info("""
        **📋 Características del formato APA:**
        - Tablas con formato académico estándar
        - Fuente Times New Roman 12pt
        - Bordes y espaciado según normas APA
        - Títulos numerados automáticamente
        - Fecha del análisis incluida
        - Compatible con editores de texto académicos
        """)

if __name__ == "__main__":
    main()
